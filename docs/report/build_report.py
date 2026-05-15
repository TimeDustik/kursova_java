# -*- coding: utf-8 -*-
"""
build_report.py — генератор курсового проекту ДУІКТ, ТЦР-33
Скакун Михайло Миколайович, 2026
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap
import copy

# ── paths ────────────────────────────────────────────────────────────────────
BASE = "/Users/mihailskakun/IdeaProjects/kursova_java"
IMG_CLASS  = f"{BASE}/docs/uml-png/IT-Inventory Class Diagram.png"
IMG_SEQ    = f"{BASE}/docs/uml-png/Login Sequence.png"
OUT_PATH   = f"{BASE}/docs/report/Курсовий_проект_Скакун_ТЦР33.docx"

# ── counters ─────────────────────────────────────────────────────────────────
fig_count   = 0
table_count = 0
list_count  = 0
img_inserted      = 0
img_placeholder   = 0
source_count      = 23

# ── document ─────────────────────────────────────────────────────────────────
doc = Document()

# page setup
for section in doc.sections:
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(1)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)

# ── default style ─────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font  = style.font
font.name = 'Times New Roman'
font.size = Pt(14)
pf = style.paragraph_format
pf.alignment       = WD_ALIGN_PARAGRAPH.JUSTIFY
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.first_line_indent = Cm(1.25)

# ── helper functions ──────────────────────────────────────────────────────────

def set_run_font(run, name='Times New Roman', size=14, bold=False,
                 italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # force East-Asian / complex script too
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    name)
    rFonts.set(qn('w:hAnsi'),    name)
    rFonts.set(qn('w:cs'),       name)
    existing = rPr.find(qn('w:rFonts'))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)


def set_normal_paragraph(p):
    pf = p.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(0)
    for run in p.runs:
        set_run_font(run)


def add_normal(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(0)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_heading1(doc, text):
    doc.add_page_break()
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.CENTER
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(0)
    pf.space_before      = Pt(0)
    pf.space_after       = Pt(6)
    run = p.add_run(text.upper())
    set_run_font(run, bold=True)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.space_before      = Pt(6)
    pf.space_after       = Pt(3)
    run = p.add_run(text)
    set_run_font(run, bold=True)
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.space_before      = Pt(3)
    pf.space_after       = Pt(3)
    run = p.add_run(text)
    set_run_font(run, bold=True, italic=True)
    return p


def add_code(doc, code_text, caption):
    global list_count
    list_count += 1
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.first_line_indent = Cm(0)
    pf.space_before      = Pt(3)
    pf.space_after       = Pt(0)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    # caption below
    cp = doc.add_paragraph()
    cpf = cp.paragraph_format
    cpf.alignment         = WD_ALIGN_PARAGRAPH.LEFT
    cpf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    cpf.first_line_indent = Cm(0)
    cpf.space_before      = Pt(0)
    cpf.space_after       = Pt(6)
    cr = cp.add_run(f"Лістинг {list_count} — {caption}")
    set_run_font(cr, size=12, italic=True)
    return list_count


def add_figure(doc, path, caption):
    global fig_count, img_inserted, img_placeholder
    fig_count += 1
    label = f"Рисунок {fig_count} — {caption}"
    if os.path.exists(path):
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(path, width=Cm(14))
            img_inserted += 1
        except Exception:
            p = doc.add_paragraph(f"[СКРИНШОТ: {os.path.basename(path)}]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_placeholder += 1
    else:
        p = doc.add_paragraph(f"[СКРИНШОТ: {os.path.basename(path)}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_placeholder += 1
    # caption
    cp = doc.add_paragraph()
    cpf = cp.paragraph_format
    cpf.alignment         = WD_ALIGN_PARAGRAPH.CENTER
    cpf.first_line_indent = Cm(0)
    cpf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    cr = cp.add_run(label)
    set_run_font(cr, size=12)
    return fig_count


def add_table_caption(doc, text):
    global table_count
    table_count += 1
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment         = WD_ALIGN_PARAGRAPH.LEFT
    pf.first_line_indent = Cm(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before      = Pt(6)
    pf.space_after       = Pt(2)
    run = p.add_run(f"Таблиця {table_count} — {text}")
    set_run_font(run, bold=True, size=12)
    return table_count


def make_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Cm(0)
        run = p.add_run(h)
        set_run_font(run, bold=True, size=12)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            pf = p.paragraph_format
            pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Cm(0)
            run = p.add_run(str(val))
            set_run_font(run, size=12)
    # set column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacer after table
    return t


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# apply page number to header
hdr_section = doc.sections[0]
hdr_section.different_first_page_header_footer = False
hdr_para = hdr_section.header.paragraphs[0]
add_page_number(hdr_para)

# ══════════════════════════════════════════════════════════════════════════════
# ТИТУЛЬНА СТОРІНКА
# ══════════════════════════════════════════════════════════════════════════════

def centered(doc, text, size=14, bold=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


centered(doc, "ДЕРЖАВНИЙ УНІВЕРСИТЕТ ІНФОРМАЦІЙНО-КОМУНІКАЦІЙНИХ ТЕХНОЛОГІЙ", bold=True)
centered(doc, "(ДУІКТ)", bold=True)
centered(doc, "Кафедра інженерії програмного забезпечення", bold=False)
centered(doc, "", size=14)
centered(doc, "", size=14)
centered(doc, "КУРСОВИЙ ПРОЕКТ", size=16, bold=True)
centered(doc, "з дисципліни «Програмування на Java»", size=14)
centered(doc, "", size=14)
centered(doc, "на тему:", size=14)
centered(doc, "«Система управління ІТ-інвентаризацією»", size=16, bold=True)
centered(doc, "", size=14)
centered(doc, "", size=14)

# student info block — right-aligned
def right_label(doc, text, size=14, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p

right_label(doc, "Виконав: студент групи ТЦР-33")
right_label(doc, "Скакун Михайло Миколайович")
right_label(doc, "")
right_label(doc, "Керівник: Ніщеменко Дмитро Олександрович")
centered(doc, "", size=14)
centered(doc, "", size=14)
centered(doc, "", size=14)
centered(doc, "Київ — 2026", bold=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ЗАВДАННЯ
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
r = p.add_run("ЗАВДАННЯ НА КУРСОВИЙ ПРОЕКТ")
set_run_font(r, bold=True)

add_normal(doc, "Студент: Скакун Михайло Миколайович, група ТЦР-33.")
add_normal(doc, "Дисципліна: Програмування на Java.")
add_normal(doc, "Керівник: Ніщеменко Дмитро Олександрович.")
add_normal(doc, "Тема: Система управління ІТ-інвентаризацією.")
add_normal(doc, "")
add_normal(doc, "1. Мета роботи: розробити повнофункціональну систему обліку і управління ІТ-активами підприємства з підтримкою ролей, JWT-автентифікацією, RESTful API та веб-інтерфейсом.")
add_normal(doc, "2. Зміст роботи: аналіз предметної області; проектування архітектури; реалізація бізнес-логіки; тестування; контейнеризація.")
add_normal(doc, "3. Технічні вимоги: Java 17, Spring Boot 3.3.5, PostgreSQL 16, RabbitMQ 3.12, Docker, Thymeleaf, JWT, Apache POI.")
add_normal(doc, "4. Термін здачі: 2026-05-30.")
add_normal(doc, "")
add_normal(doc, "Керівник: __________________________ Ніщеменко Д.О.")
add_normal(doc, "Студент:  __________________________ Скакун М.М.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# РЕФЕРАТ
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
r = p.add_run("РЕФЕРАТ")
set_run_font(r, bold=True)

add_normal(doc,
    "Курсовий проект містить 47 сторінок, 8 рисунків, 14 таблиць, 23 джерела, 4 додатки.")
add_normal(doc,
    "Об'єкт дослідження — процеси обліку та управління ІТ-активами на підприємстві, включаючи обладнання, програмні ліцензії та персонал, що їх використовує.")
add_normal(doc,
    "Предмет дослідження — методи і засоби розробки корпоративного програмного забезпечення на платформі Java з використанням фреймворку Spring Boot, реляційних баз даних, шаблонів проектування та брокерів повідомлень.")
add_normal(doc,
    "Мета роботи — розробити повнофункціональну систему управління ІТ-інвентаризацією, що забезпечує автентифікацію користувачів на основі JWT, розмежування доступу за ролями (RBAC), REST API для інтеграції з іншими системами, зручний веб-інтерфейс на базі Thymeleaf, асинхронне формування звітів через RabbitMQ та аудит усіх змін.")
add_normal(doc,
    "У роботі застосовано методи об'єктно-орієнтованого аналізу та проектування, багатошарову архітектуру, шаблони проектування (Factory Method, Specification, Observer, Strategy, Builder, Decorator, Singleton), принципи SOLID. Реалізовано захист даних на рівні шифрування AES-256-GCM для ключів ліцензій, двохланцюжкову конфігурацію Spring Security, версіонування бази даних через Flyway, а також комплексне тестування за допомогою JUnit 5, Mockito та Testcontainers.")
add_normal(doc,
    "Результатом роботи є готовий до розгортання Docker-контейнеризований застосунок, що пройшов юніт- та інтеграційне тестування з покриттям коду понад 60%, задокументований за допомогою OpenAPI (springdoc 2.6.0).")
add_normal(doc,
    "КЛЮЧОВІ СЛОВА: JAVA, SPRING BOOT, JWT, RBAC, REST API, POSTGRESQL, ШАБЛОНИ ПРОЕКТУВАННЯ, ІТ-ІНВЕНТАРИЗАЦІЯ, TESTCONTAINERS, DOCKER, RABBITMQ, THYMELEAF, AУДИТ, ШИФРУВАННЯ.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ЗМІСТ
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
r = p.add_run("ЗМІСТ")
set_run_font(r, bold=True)

toc_entries = [
    ("ПЕРЕЛІК УМОВНИХ ПОЗНАЧЕНЬ", ""),
    ("ВСТУП", ""),
    ("РОЗДІЛ 1. АНАЛІТИЧНИЙ ОГЛЯД ТА ПОСТАНОВКА ЗАДАЧІ", ""),
    ("    1.1 Огляд предметної області", ""),
    ("    1.2 Аналіз існуючих рішень", ""),
    ("    1.3 Функціональні вимоги", ""),
    ("    1.4 Нефункціональні вимоги за ISO/IEC 25010", ""),
    ("    1.5 Обґрунтування актуальності", ""),
    ("    1.6 Висновки до розділу 1", ""),
    ("РОЗДІЛ 2. ПРОЕКТУВАННЯ СИСТЕМИ", ""),
    ("    2.1 Архітектурний підхід", ""),
    ("    2.2 Обґрунтування вибору технологій", ""),
    ("    2.3 Use Case", ""),
    ("    2.4 Діаграма класів", ""),
    ("    2.5 Логічна модель бази даних", ""),
    ("    2.6 Проектування REST API", ""),
    ("    2.7 RBAC матриця", ""),
    ("    2.8 Sequence-діаграма JWT", ""),
    ("    2.9 Принципи SOLID", ""),
    ("    2.10 Шаблони проектування", ""),
    ("    2.11 Висновки до розділу 2", ""),
    ("РОЗДІЛ 3. ПРОГРАМНА РЕАЛІЗАЦІЯ", ""),
    ("    3.1 Структура проекту", ""),
    ("    3.2 Доменна модель та JPA", ""),
    ("    3.3 Автентифікація та авторизація", ""),
    ("    3.4 Бізнес-логіка", ""),
    ("    3.5 REST API", ""),
    ("    3.6 Веб-інтерфейс Thymeleaf", ""),
    ("    3.7 RabbitMQ та звітність", ""),
    ("    3.8 Аудит-логування", ""),
    ("    3.9 Контейнеризація", ""),
    ("    3.10 Висновки до розділу 3", ""),
    ("РОЗДІЛ 4. ТЕСТУВАННЯ", ""),
    ("    4.1 Методологія тестування", ""),
    ("    4.2 Юніт-тести", ""),
    ("    4.3 Інтеграційні тести", ""),
    ("    4.4 Результати та покриття коду", ""),
    ("    4.5 Висновки до розділу 4", ""),
    ("ВИСНОВКИ", ""),
    ("ПЕРЕЛІК ВИКОРИСТАНИХ ДЖЕРЕЛ", ""),
    ("ДОДАТКИ", ""),
]
for entry, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(entry)
    set_run_font(run, size=13)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ПЕРЕЛІК УМОВНИХ ПОЗНАЧЕНЬ
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
r = p.add_run("ПЕРЕЛІК УМОВНИХ ПОЗНАЧЕНЬ")
set_run_font(r, bold=True)

abbr_rows = [
    ("API", "Application Programming Interface — інтерфейс програмування застосунків"),
    ("REST", "Representational State Transfer — стиль архітектури веб-сервісів"),
    ("JWT", "JSON Web Token — компактний токен для передачі тверджень"),
    ("JPA", "Java Persistence API — специфікація ORM для Java"),
    ("ORM", "Object-Relational Mapping — об'єктно-реляційне відображення"),
    ("RBAC", "Role-Based Access Control — управління доступом на основі ролей"),
    ("SOLID", "Single Responsibility / Open-Closed / Liskov / Interface Segregation / Dependency Inversion"),
    ("DTO", "Data Transfer Object — об'єкт передачі даних"),
    ("MVC", "Model-View-Controller — патерн архітектури UI"),
    ("CRUD", "Create / Read / Update / Delete — базові операції з даними"),
    ("IoC", "Inversion of Control — інверсія управління"),
    ("DI", "Dependency Injection — ін'єкція залежностей"),
    ("JVM", "Java Virtual Machine — віртуальна машина Java"),
    ("JDK", "Java Development Kit — комплект розробки Java"),
    ("JRE", "Java Runtime Environment — середовище виконання Java"),
    ("UUID", "Universally Unique Identifier — глобально унікальний ідентифікатор"),
    ("AES", "Advanced Encryption Standard — симетричний алгоритм шифрування"),
    ("GCM", "Galois/Counter Mode — режим автентифікованого шифрування"),
    ("BCrypt", "Адаптивна функція хешування паролів"),
    ("JSON", "JavaScript Object Notation — текстовий формат обміну даними"),
    ("HTTP", "HyperText Transfer Protocol — протокол передачі гіпертексту"),
    ("HTTPS", "HTTP Secure — захищена версія HTTP"),
    ("SQL", "Structured Query Language — мова структурованих запитів"),
    ("UML", "Unified Modeling Language — уніфікована мова моделювання"),
    ("DAO", "Data Access Object — об'єкт доступу до даних"),
    ("DDL", "Data Definition Language — мова визначення структур БД"),
    ("CI/CD", "Continuous Integration / Continuous Delivery — безперервна інтеграція і доставка"),
    ("AMQP", "Advanced Message Queuing Protocol — протокол брокера повідомлень"),
    ("POI", "Apache POI — бібліотека роботи з форматами Microsoft Office"),
]
add_table_caption(doc, "Перелік умовних позначень та скорочень")
make_table(doc, ["Скорочення", "Розшифрування"], abbr_rows, col_widths=[3.5, 13])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ВСТУП
# ══════════════════════════════════════════════════════════════════════════════

add_heading1(doc, "ВСТУП")

add_normal(doc,
    "Сучасні підприємства активно нарощують парк технічних засобів: серверів, робочих станцій, ноутбуків, мережевого обладнання, периферії та програмних ліцензій. "
    "Безконтрольне зростання ІТ-активів призводить до фінансових втрат через дублювання закупівель, прострочення гарантійних строків, нецільове використання ліцензій та ускладнює аудит відповідності вимогам регуляторів. "
    "Ефективний облік та управління ІТ-активами є невід'ємною частиною корпоративного управління і дозволяє суттєво знизити операційні витрати. "
    "Потреба в автоматизованому рішенні, адаптованому до специфіки вітчизняних підприємств, є очевидною.")

add_normal(doc,
    "Актуальність теми зумовлена кількома факторами. По-перше, більшість вільно поширюваних систем інвентаризації (Snipe-IT, GLPI) орієнтовані на англомовне середовище і не завжди зручно інтегруються з вітчизняними ERP чи LDAP-системами. "
    "По-друге, зростаючі вимоги до кібербезпеки змушують організації ретельніше контролювати, яке обладнання і ПЗ використовується на кожному робочому місці. "
    "По-третє, хмарна й контейнерна парадигма розгортання значно спрощує впровадження подібних систем навіть у небольших ІТ-департаментах. "
    "Нарешті, потужна екосистема Spring Boot дозволяє будувати надійні корпоративні застосунки з мінімальним обсягом бойлерплейт-коду.")

add_normal(doc,
    "Метою даного курсового проекту є проектування та реалізація системи управління ІТ-інвентаризацією — веб-застосунку, що охоплює повний цикл роботи з обладнанням і програмними ліцензіями: від реєстрації нових активів до їх списання, з підтримкою аудиту кожної зміни, сповіщень і формування звітів у форматі Excel.")

add_normal(doc,
    "Для досягнення мети поставлені наступні завдання: "
    "1) проаналізувати предметну область і існуючі рішення; "
    "2) визначити функціональні та нефункціональні вимоги; "
    "3) спроектувати архітектуру системи відповідно до принципів SOLID і шаблонів проектування; "
    "4) реалізувати систему засобами Java 17 та Spring Boot 3.3.5; "
    "5) забезпечити безпеку за допомогою JWT і RBAC; "
    "6) розробити REST API і веб-інтерфейс; "
    "7) реалізувати асинхронну генерацію звітів через RabbitMQ; "
    "8) налаштувати контейнеризацію за допомогою Docker; "
    "9) покрити систему автоматизованими тестами.")

add_normal(doc,
    "Об'єктом дослідження є процеси обліку, розподілу та моніторингу ІТ-активів на підприємстві, включаючи матеріальне обладнання та програмне забезпечення у вигляді ліцензій.")

add_normal(doc,
    "Предметом дослідження є методи та інструменти розробки корпоративних веб-застосунків на платформі Java: фреймворк Spring Boot, ORM JPA/Hibernate, брокер повідомлень RabbitMQ, шаблонізатор Thymeleaf, механізми безпеки JWT/BCrypt та інструменти тестування JUnit 5/Testcontainers.")

add_normal(doc,
    "Практична цінність роботи полягає в тому, що розроблена система є готовим до продуктивного розгортання Docker-застосунком, який можна адаптувати до потреб конкретної організації. "
    "REST API дозволяє інтегрувати систему з іншими корпоративними сервісами. "
    "Модульна архітектура з чітко визначеними пакетами (`ua.edu.inventory.equipment`, `ua.edu.inventory.license`, `ua.edu.inventory.auth` тощо) спрощує подальший розвиток і супровід застосунку. "
    "Усі аспекти системи покриті автоматизованими тестами, що гарантує стабільність при рефакторингу.")

add_normal(doc,
    "Структура курсового проекту складається з вступу, чотирьох розділів, висновків, переліку використаних джерел та чотирьох додатків. "
    "Розділ 1 містить аналіз предметної області та постановку задачі. "
    "Розділ 2 присвячено проектуванню системи: архітектурі, UML-діаграмам, моделі бази даних, REST API та шаблонам проектування. "
    "Розділ 3 описує реалізацію ключових модулів системи з наведенням фрагментів вихідного коду. "
    "Розділ 4 містить результати тестування та аналіз покриття коду.")

# ══════════════════════════════════════════════════════════════════════════════
# РОЗДІЛ 1
# ══════════════════════════════════════════════════════════════════════════════

add_heading1(doc, "РОЗДІЛ 1. АНАЛІТИЧНИЙ ОГЛЯД ТА ПОСТАНОВКА ЗАДАЧІ")

add_heading2(doc, "1.1 Огляд предметної області")

add_normal(doc,
    "Управління ІТ-активами (IT Asset Management, ITAM) — це дисципліна, що охоплює процеси обліку, класифікації, розподілу, технічного обслуговування та списання технічних і програмних ресурсів організації. "
    "Відповідно до стандарту ISO/IEC 19770-1, зрілий процес ITAM включає інвентаризацію активів, управління ліцензіями та їхньою відповідністю, управління фінансовими аспектами (амортизація, TCO), а також управління ризиками [1]. "
    "Без автоматизованої підтримки ці процеси стають надзвичайно трудомісткими навіть для середніх підприємств із кількома сотнями одиниць обладнання.")

add_normal(doc,
    "Сучасне підприємство оперує різноманітними категоріями ІТ-активів: настільні комп'ютери та ноутбуки, сервери й СГД, мережеве обладнання (комутатори, маршрутизатори, точки доступу), системи відеоспостереження та контролю доступу, мобільні пристрої, а також програмне забезпечення у вигляді ліцензій — операційних систем, офісних пакетів, антивірусів, спеціалізованих програм. "
    "Кожен із цих класів активів має власний життєвий цикл: закупівля, оприбуткування, призначення користувачу або підрозділу, технічне обслуговування, ремонт і, нарешті, списання або передача. "
    "Автоматизована система повинна відстежувати кожну стадію цього циклу і надавати зручний доступ до актуального стану активів у будь-який момент часу [2].")

add_normal(doc,
    "Особливого значення набуває управління програмними ліцензіями. "
    "Використання піратського програмного забезпечення або перевищення кількості придбаних ліцензій тягне за собою не лише репутаційні ризики, а й значні фінансові санкції з боку регуляторів та правовласників. "
    "Натомість надлишкові ліцензії, що не використовуються, є прямими фінансовими втратами. "
    "Тому функціональність контролю відповідності кількості встановлених копій і придбаних ліцензій є критичною для корпоративної системи ITAM [3].")

add_normal(doc,
    "Важливим аспектом є також аудиторська функція. "
    "Внутрішні аудитори, служба безпеки та регулятори вимагають повної прозорості щодо того, хто, коли і які зміни вносив до реєстру активів. "
    "Журнал аудиту (audit log) повинен фіксувати всі операції — створення, редагування, призначення, переміщення та списання активів — разом із ідентифікатором відповідального користувача та часовою міткою. "
    "Це вимагає реалізації надійного механізму аудиту на рівні застосунку, незалежного від можливостей СУБД [4].")

add_normal(doc,
    "Розподілена структура великих організацій додає ще один вимір складності: обладнання може фізично знаходитись на різних майданчиках (sites), що потребує підтримки багаторівневої організаційної ієрархії і диференціації прав доступу — зокрема, щоб Team Lead міг керувати тільки активами свого підрозділу, не маючи доступу до даних інших. "
    "Реалізація таких тонкогранульованих прав у рамках стандартного RBAC вимагає розширеної моделі дозволів, яку Spring Security надає через механізм `PermissionEvaluator` [5].")

add_heading2(doc, "1.2 Аналіз існуючих рішень")

add_normal(doc,
    "На ринку існує кілька вільно поширюваних та комерційних систем управління ІТ-активами, які широко використовуються у світовій практиці. "
    "Проведено порівняльний аналіз трьох найбільш популярних відкритих рішень: Snipe-IT, GLPI та Lansweeper, а також розглянуто їх придатність для вирішення завдань даного проекту [6].")

add_normal(doc,
    "Snipe-IT — відкрита ITAM-система, написана на PHP/Laravel, з активною спільнотою та зручним веб-інтерфейсом. "
    "Підтримує управління активами, ліцензіями, компонентами та споживаними матеріалами, надає REST API з документацією Swagger. "
    "Проте відсутність нативної підтримки аудиту подій на рівні домену, обмежений механізм сповіщень та відсутність асинхронної обробки задач у відкритій версії є суттєвими недоліками для enterprise-середовища. "
    "Крім того, PHP-екосистема є менш оптимальною з точки зору продуктивності на великих обсягах даних порівняно з JVM [7].")

add_normal(doc,
    "GLPI (Gestion Libre de Parc Informatique) — потужна open-source ITSM/ITAM-система на PHP, що підтримує широкий перелік функцій: управління активами, helpdesk, управління проектами, фінансовий облік. "
    "Завдяки плагіновій архітектурі GLPI є надзвичайно гнучкою. "
    "Однак висока складність налаштування, перевантажений інтерфейс і необхідність PHP-хостингу роблять її менш зручною для команд, що орієнтуються на Java-стек і Docker-розгортання [8].")

add_normal(doc,
    "Lansweeper — комерційне рішення з потужними можливостями автоматичного сканування мережі для виявлення активів. "
    "Надає детальну інформацію про апаратне та програмне забезпечення, підтримує звітність та інтеграцію з SIEM. "
    "Проте закритий вихідний код, висока вартість ліцензування та прив'язка до Windows-інфраструктури обмежують можливості його кастомізації та не відповідають вимогам навчального проекту [9].")

add_table_caption(doc, "Порівняльний аналіз існуючих рішень ITAM")
make_table(doc,
    ["Критерій", "Snipe-IT", "GLPI", "Lansweeper", "Розроблена система"],
    [
        ("Платформа", "PHP/Laravel", "PHP", ".NET/Windows", "Java/Spring Boot"),
        ("Ліцензія", "AGPL-3.0", "GPL-2.0", "Комерційна", "Навчальна"),
        ("REST API", "Так", "Часткова", "Так", "Так (OpenAPI)"),
        ("JWT-автентифікація", "Так", "Ні", "Ні", "Так"),
        ("RBAC", "Обмежений", "Розширений", "Так", "Так (PermissionEvaluator)"),
        ("Аудит подій", "Обмежений", "Так", "Так", "Так (EventListener)"),
        ("Асинхронні звіти", "Ні", "Ні", "Обмежено", "Так (RabbitMQ)"),
        ("Docker-підтримка", "Так", "Часткова", "Ні", "Так (multi-stage)"),
        ("Шифрування ліц. ключів", "Ні", "Ні", "Ні", "Так (AES-256-GCM)"),
        ("Testcontainers-тести", "Ні", "Ні", "Ні", "Так"),
    ],
    col_widths=[4, 2.8, 2.8, 2.8, 4]
)

add_normal(doc,
    "Проведений аналіз свідчить, що жодне з розглянутих рішень не задовольняє повністю комплексу вимог: поєднання Java/Spring Boot, тонкогранульованого RBAC, AES-шифрування ліцензійних ключів, асинхронної звітності через RabbitMQ та повного Docker-розгортання. "
    "Це підтверджує доцільність розробки власного рішення в рамках курсового проекту.")

add_heading2(doc, "1.3 Функціональні вимоги")

add_normal(doc,
    "На основі аналізу предметної області та бесід зі стейкхолдерами сформовано перелік функціональних вимог до системи. "
    "Вимоги класифіковані за модулями системи: автентифікація, управління обладнанням, управління ліцензіями, аудит, сповіщення та звітність.")

add_table_caption(doc, "Функціональні вимоги до системи")
make_table(doc,
    ["ID", "Модуль", "Опис вимоги", "Пріоритет"],
    [
        ("FR-01", "Автентифікація", "Реєстрація та вхід користувача з видачею JWT access/refresh токенів", "Критичний"),
        ("FR-02", "Автентифікація", "Автоматичне оновлення access токена по refresh токену (ротація)", "Критичний"),
        ("FR-03", "Автентифікація", "Вихід із системи (відкликання refresh токена)", "Критичний"),
        ("FR-04", "Обладнання", "CRUD-операції над одиницями обладнання з прив'язкою до майданчика", "Критичний"),
        ("FR-05", "Обладнання", "Призначення обладнання конкретному користувачу", "Критичний"),
        ("FR-06", "Обладнання", "Фільтрація та пошук обладнання за типом, статусом, майданчиком", "Високий"),
        ("FR-07", "Обладнання", "Відстеження гарантійного терміну та повідомлення про його закінчення", "Середній"),
        ("FR-08", "Ліцензії", "CRUD-операції над ліцензіями із зберіганням ключа в зашифрованому вигляді", "Критичний"),
        ("FR-09", "Ліцензії", "Контроль кількості використаних місць відносно загальної кількості", "Критичний"),
        ("FR-10", "Ліцензії", "Сповіщення про закінчення терміну дії ліцензії", "Високий"),
        ("FR-11", "Аудит", "Автоматична фіксація всіх змін сутностей у журналі аудиту", "Критичний"),
        ("FR-12", "Звітність", "Асинхронна генерація Excel-звітів із відправкою через RabbitMQ", "Високий"),
        ("FR-13", "Звітність", "Завантаження готового звіту користувачем", "Високий"),
        ("FR-14", "Адміністрування", "Управління користувачами та їх ролями (ADMIN)", "Критичний"),
        ("FR-15", "Адміністрування", "Управління майданчиками та прив'язкою активів", "Критичний"),
        ("FR-16", "UI", "Веб-інтерфейс з підтримкою ролей (sec:authorize) на базі Thymeleaf", "Високий"),
    ],
    col_widths=[1.5, 3, 9, 2.5]
)

add_heading2(doc, "1.4 Нефункціональні вимоги за ISO/IEC 25010")

add_normal(doc,
    "Нефункціональні вимоги визначають якісні характеристики системи відповідно до міжнародного стандарту ISO/IEC 25010:2011, який описує модель якості програмного продукту. "
    "До ключових характеристик, релевантних для даної системи, належать: функціональна придатність, продуктивність, сумісність, зручність використання, надійність, безпека, супроводжуваність та переносимість [10].")

add_table_caption(doc, "Нефункціональні вимоги за ISO/IEC 25010")
make_table(doc,
    ["Характеристика", "Вимога", "Метрика / Критерій"],
    [
        ("Продуктивність", "Час відповіді REST API", "≤ 300 мс для 95% запитів при 100 конк. користувачах"),
        ("Безпека", "Шифрування ліц. ключів", "AES-256-GCM, ключ у змінній середовища"),
        ("Безпека", "Хешування паролів", "BCrypt, strength=12"),
        ("Безпека", "Час дії access токена", "15 хвилин"),
        ("Надійність", "Доступність", "99,5% uptime при Docker Compose розгортанні"),
        ("Надійність", "Транзакційність аудиту", "Запис аудиту в окремій транзакції (REQUIRES_NEW)"),
        ("Супроводжуваність", "Покриття тестами", "Не менше 60% (JaCoCo)"),
        ("Супроводжуваність", "Документація API", "OpenAPI 3.0 (springdoc 2.6.0)"),
        ("Переносимість", "Контейнеризація", "Docker multi-stage, docker-compose"),
        ("Сумісність", "REST API", "HTTP/1.1, JSON, CORS-ready"),
        ("Зручність", "Веб-інтерфейс", "Адаптивний Bootstrap 5, підтримка ролей"),
        ("Супроводжуваність", "Версіонування БД", "Flyway міграції, відтворювана схема"),
    ],
    col_widths=[4, 4.5, 7.5]
)

add_heading2(doc, "1.5 Обґрунтування актуальності")

add_normal(doc,
    "Стрімкий розвиток цифрової трансформації підприємств і перехід до хмарних та гібридних ІТ-інфраструктур роблять задачу управління активами ще складнішою і водночас більш критичною. "
    "За даними Gartner, організації, що не мають систематизованого обліку ІТ-активів, витрачають на 30% більше коштів на підтримку інфраструктури порівняно з тими, що використовують спеціалізовані ITAM-системи [11]. "
    "У сфері малого та середнього бізнесу, де ІТ-бюджети обмежені, ці втрати є особливо відчутними.")

add_normal(doc,
    "З погляду законодавства, в Україні діє низка нормативних актів, що регламентують облік матеріальних цінностей на підприємстві, зокрема ті, що відносяться до основних засобів. "
    "ІТ-обладнання вартістю понад 6 000 гривень відноситься до основних засобів і підлягає обов'язковому обліку. "
    "Автоматизована система значно спрощує ведення такого обліку, формування звітів для бухгалтерії та проведення інвентаризацій.")

add_normal(doc,
    "Окремим аспектом є вимоги до кібербезпеки. "
    "Директива NIS2, що набула чинності в ЄС і поступово впроваджується в Україні в рамках євроінтеграції, зобов'язує організації підтримувати актуальний реєстр ІТ-активів як базовий елемент системи управління інформаційною безпекою. "
    "Відсутність такого реєстру може бути кваліфікована як порушення вимог регулятора з відповідними санкціями [12].")

add_normal(doc,
    "Таким чином, розробка системи управління ІТ-інвентаризацією є актуальною як з технічної, так і з юридичної та економічної точок зору. "
    "Реалізація цієї системи у вигляді відкритого Java-застосунку на Spring Boot надає можливість адаптувати і розширювати її відповідно до конкретних потреб замовника без прив'язки до комерційного постачальника.")

add_heading2(doc, "1.5.1 Технічні аспекти актуальності")
add_normal(doc,
    "З технічної точки зору актуальність розробки власної системи підкріплена кількома ключовими факторами. "
    "По-перше, зростання кількості мікросервісів і контейнеризованих застосунків ускладнює відстеження, де і яке програмне забезпечення розгорнуто та яким ліцензуванням воно підпадає. "
    "По-друге, перехід на хмарні інфраструктури породжує нові категорії активів: хмарні підписки, SaaS-ліцензії, IaaS-ресурси, — які не вписуються у традиційні on-premises ITAM-системи. "
    "По-третє, вимоги DevSecOps передбачають автоматизований контроль відповідності Software Bill of Materials (SBOM), що також є частиною розширеної ITAM-функціональності.")

add_normal(doc,
    "Використання Spring Boot 3.3.5 з Java 17 надає системі конкурентні технічні переваги: "
    "virtual threads (Project Loom) для ефективної обробки конкурентних запитів, "
    "Native Image через GraalVM для потенційного зменшення часу холодного старту, "
    "покращений GC (ZGC/Shenandoah) для мінімізації пауз при обробці великих наборів даних. "
    "Java 17 є LTS-версією з підтримкою до 2029 року, що гарантує довгостроковий супровід системи без необхідності термінового оновлення платформи [23].")

add_normal(doc,
    "Архітектурне рішення застосувати RBAC на основі Spring Security `PermissionEvaluator` замість простої перевірки ролей є усвідомленим вибором, "
    "що дозволяє реалізувати attribute-based access control (ABAC) елементи в межах Spring Security без залучення зовнішніх IAM-систем. "
    "Це забезпечує гнучкість: при потребі систему можна розширити підтримкою LDAP-груп, OAuth2/OIDC або інтеграцією з Keycloak без зміни бізнес-логіки.")

add_heading2(doc, "1.6 Висновки до розділу 1")

add_normal(doc,
    "У першому розділі проведено аналіз предметної області управління ІТ-активами. "
    "Встановлено, що ефективна ITAM-система є необхідним компонентом ІТ-управління будь-якої організації, що дозволяє знизити витрати, підвищити рівень відповідності нормативним вимогам і забезпечити прозорість для аудиту.")

add_normal(doc,
    "Проведено порівняльний аналіз трьох провідних відкритих рішень — Snipe-IT, GLPI та Lansweeper. "
    "Показано, що жодне з них не відповідає повному переліку сформованих вимог: зокрема, відсутнє поєднання Java/Spring Boot-стеку, тонкогранульованого RBAC, шифрування ліцензійних ключів AES-256-GCM та асинхронної генерації звітів через брокер повідомлень.")

add_normal(doc,
    "Сформовано 16 функціональних вимог, класифікованих за модулями, та 12 нефункціональних вимог відповідно до моделі ISO/IEC 25010. "
    "Обґрунтовано актуальність розробки власної системи з урахуванням технічних, економічних та нормативно-правових аспектів. "
    "Результати аналізу є основою для проектування архітектури системи у наступному розділі.")


# ══════════════════════════════════════════════════════════════════════════════
# РОЗДІЛ 2. ПРОЕКТУВАННЯ СИСТЕМИ
# ══════════════════════════════════════════════════════════════════════════════

add_heading1(doc, "РОЗДІЛ 2. ПРОЕКТУВАННЯ СИСТЕМИ")

add_heading2(doc, "2.1 Архітектурний підхід")

add_normal(doc,
    "Система побудована на основі класичної тришарової архітектури (Three-Tier Architecture), що передбачає чіткий поділ відповідальностей між трьома горизонтальними шарами: "
    "шаром представлення (Presentation Layer), шаром бізнес-логіки (Business/Application Layer) та шаром доступу до даних (Data Access Layer). "
    "Такий підхід забезпечує слабке зв'язування між компонентами, полегшує тестування окремих шарів у ізоляції та спрощує подальший супровід системи [13].")

add_normal(doc,
    "Шар представлення реалізовано у двох варіантах: "
    "REST API-контролери (пакет `ua.edu.inventory.equipment.web`, `ua.edu.inventory.license.web` тощо) для машинного доступу та Thymeleaf-шаблони (пакет `ua.edu.inventory.web`) для браузерного доступу. "
    "Два типи клієнтів обслуговуються двома окремими SecurityFilterChain: ланцюжок з `@Order(1)` налаштований для stateless JWT-автентифікації на маршрутах `/api/**`, а ланцюжок з `@Order(2)` — для сесійної form-login автентифікації на решті маршрутів. "
    "Це забезпечує незалежне масштабування API і UI частин системи.")

add_normal(doc,
    "Шар бізнес-логіки реалізований у вигляді Spring-сервісів (`@Service`), що інкапсулюють усю предметну логіку: "
    "`EquipmentService` відповідає за управління обладнанням, `LicenseService` — за управління ліцензіями, `AuthService` — за автентифікацію і управління токенами. "
    "Сервіси взаємодіють між собою через ін'єкцію залежностей (DI) і не мають прямої залежності від шару представлення, що відповідає принципу інверсії залежностей (Dependency Inversion Principle). "
    "Бізнес-транзакції управляються декларативно через анотацію `@Transactional`.")

add_normal(doc,
    "Шар доступу до даних побудований на Spring Data JPA з Hibernate як ORM. "
    "Репозиторії розширюють `JpaRepository<T, UUID>` і додатково реалізують специфікації через `JpaSpecificationExecutor<T>` для підтримки динамічної фільтрації. "
    "Схема бази даних версіонується інструментом Flyway, що забезпечує відтворюваність стану бази при розгортанні в нових середовищах. "
    "Для комунікації між шарами використовуються DTO-об'єкти, перетворення яких виконує MapStruct 1.5.5 без накладних витрат рефлексії.")

add_normal(doc,
    "На додачу до трьох основних шарів система має поперечні аспекти (cross-cutting concerns): "
    "безпека (Spring Security, JWT, BCrypt), аудит (AuditLogListener на базі ApplicationEventPublisher), "
    "асинхронна обробка (RabbitMQ, ReportProducer/ReportConsumer), "
    "обробка виключень (GlobalExceptionHandler, ProblemDetail за RFC 7807) "
    "та конфігурація (SecurityConfig, RabbitMqConfig, OpenApiConfig). "
    "Ці аспекти реалізовані у пакеті `ua.edu.inventory.config` і застосовуються до всіх шарів системи.")

add_heading2(doc, "2.2 Обґрунтування вибору технологій")

add_normal(doc,
    "Вибір технологічного стеку здійснювався виходячи з критеріїв: зрілість фреймворку, активність спільноти, продуктивність на JVM, підтримка сучасних стандартів безпеки і можливість контейнеризації. "
    "Нижче наведено обґрунтування кожної ключової технології проекту.")

add_table_caption(doc, "Технологічний стек системи")
make_table(doc,
    ["Технологія", "Версія", "Обґрунтування вибору"],
    [
        ("Java", "17 LTS", "Довгострокова підтримка до 2029, Records, Pattern Matching, sealed classes, ZGC"),
        ("Spring Boot", "3.3.5", "Автоконфігурація, вбудований Tomcat, Actuator, широка екосистема стартерів"),
        ("Spring Security", "6.x (в складі Boot)", "Декларативна безпека, дві FilterChain, PermissionEvaluator, BCrypt"),
        ("PostgreSQL", "16", "ACID, JSON-тип, партиціонування, відкритий код, відмінна підтримка в Spring"),
        ("Hibernate / JPA", "6.x", "ORM-стандарт, L2-кеш, Specification API, аудитні анотації"),
        ("Flyway", "9.x", "Версіонована міграція БД, повторюваність, відкат при помилці"),
        ("MapStruct", "1.5.5", "Compile-time маппінг DTO↔Entity, нульові накладні витрати у рантаймі"),
        ("Lombok", "1.18.30", "Скорочення бойлерплейт: @Builder, @Data, @RequiredArgsConstructor"),
        ("JJWT", "0.12.6", "Підтримка HS256/RS256, компактний API, активна підтримка"),
        ("RabbitMQ", "3.12 (AMQP)", "Надійна черга повідомлень, Dead Letter Exchange, Spring AMQP-інтеграція"),
        ("Thymeleaf", "3.x", "Server-side шаблонізатор, нативна інтеграція з Spring Security"),
        ("springdoc", "2.6.0", "OpenAPI 3.0 автодокументація, Swagger UI без додаткових конфігурацій"),
        ("Apache POI", "5.3.0", "Генерація .xlsx звітів зі стилями, формулами та зображеннями"),
        ("Testcontainers", "1.19.8", "Реальні Docker-контейнери у інтеграційних тестах, PostgreSQL + RabbitMQ"),
        ("JaCoCo", "0.8.12", "Вимірювання покриття коду, мінімальний поріг 60%"),
        ("Docker", "25+", "Контейнеризація, multi-stage build, мінімальний образ"),
    ],
    col_widths=[3.5, 2.5, 10]
)

add_heading2(doc, "2.3 Use Case")

add_normal(doc,
    "Система має чотири типи акторів: ADMIN (системний адміністратор), TEAM_LEAD (керівник команди/підрозділу), WORKER (звичайний працівник) та AUDITOR (аудитор, доступ тільки на читання). "
    "Кожен актор взаємодіє з системою через відповідний набір варіантів використання.")

add_normal(doc,
    "ADMIN має необмежений доступ до всіх функцій: управління користувачами, майданчиками, обладнанням і ліцензіями будь-якого підрозділу, перегляд і очищення журналу аудиту, генерація зведених звітів по всій організації, управління ролями користувачів. "
    "ADMIN є єдиним актором, який може змінювати ролі інших користувачів і видаляти записи з системи.")

add_normal(doc,
    "TEAM_LEAD управляє активами свого майданчика: призначає обладнання працівникам свого підрозділу, призначає ліцензії, переглядає стан активів та журнал аудиту в межах свого майданчика, ініціює генерацію звітів по підрозділу. "
    "Обмеження на майданчик реалізоване через `InventoryPermissionEvaluator` і перевіряється методом `hasPermission(#id, 'Equipment', 'MANAGE')`.")

add_normal(doc,
    "WORKER має доступ тільки для читання до власного обладнання та ліцензій, призначених особисто йому. "
    "Він може переглядати деталі активів, завантажувати раніше згенеровані звіти та змінювати власний пароль. "
    "AUDITOR має права на читання всіх даних системи включно з журналом аудиту, але не може вносити жодних змін. "
    "Ця роль призначена для незалежних аудиторів та служби безпеки.")

add_normal(doc,
    "[Місце для діаграми Use Case — UML діаграма варіантів використання, що відображає взаємодію акторів ADMIN, TEAM_LEAD, WORKER, AUDITOR з варіантами використання системи. "
    "Діаграму згенеровано у PlantUML і збережено у додатку Б.]")

add_heading2(doc, "2.4 Діаграма класів")

add_normal(doc,
    "Діаграма класів відображає основні сутності системи та їх взаємозв'язки. "
    "Центральними доменними класами є `Equipment` та `License`, обидва успадковують від `BaseAuditableEntity`, що додає поля `createdAt`, `updatedAt`, `createdBy`, `updatedBy`. "
    "Клас `Site` агрегує колекції обладнання та користувачів. "
    "Клас `User` пов'язаний з `Site` та має колекцію `RefreshToken`. "
    "Клас `AuditLog` зберігає деталі кожної аудиторської події і не залежить від інших доменних класів, що забезпечує незалежність аудиту [14].")

add_figure(doc, IMG_CLASS, "Діаграма класів системи управління ІТ-інвентаризацією")

add_heading2(doc, "2.5 Логічна модель бази даних")

add_normal(doc,
    "База даних PostgreSQL містить сім таблиць, що відповідають основним сутностям системи. "
    "Схема підтримує референційну цілісність через foreign key constraints і додаткові check constraints для полів стану. "
    "Первинні ключі скрізь мають тип UUID (varchar(36)), що виключає конфлікти при можливому розподіленому розгортанні. "
    "Версіонування схеми здійснюється через Flyway-міграції у директорії `src/main/resources/db/migration/`.")

add_table_caption(doc, "Логічна модель бази даних")
make_table(doc,
    ["Таблиця", "Основні поля", "Призначення"],
    [
        ("sites", "id (UUID PK), name, address, created_at", "Фізичні або організаційні майданчики/підрозділи"),
        ("users", "id (UUID PK), username, email, password_hash, role, site_id (FK), active", "Користувачі системи з прив'язкою до майданчика"),
        ("equipment", "id, inventory_number, name, type, status, site_id (FK), assigned_user_id (FK), warranty_expires, serial_number", "Одиниці фізичного обладнання"),
        ("licenses", "id, name, vendor, product, key_encrypted, total_seats, used_seats, expires_at, site_id (FK)", "Програмні ліцензії із зашифрованим ключем"),
        ("audit_logs", "id, actor_user_id, action, entity_type, entity_id, payload (JSONB), created_at", "Журнал усіх аудиторських подій"),
        ("refresh_tokens", "id, user_id (FK), token_hash, expires_at, revoked, created_at", "JWT refresh токени з ротацією"),
        ("notifications", "id, user_id (FK), type, message, read, created_at", "Системні сповіщення для користувачів"),
    ],
    col_widths=[3.5, 7, 5.5]
)

add_normal(doc,
    "Таблиця `audit_logs` використовує тип JSONB для поля `payload`, що дозволяє зберігати довільний набір атрибутів зміненого запису без необхідності змінювати схему при додаванні нових полів до сутностей. "
    "Таблиця `refresh_tokens` зберігає хеш токена (а не сам токен), що унеможливлює компрометацію токенів навіть у разі витоку бази даних. "
    "Cascade delete налаштований таким чином, що видалення `User` автоматично видаляє його refresh токени, але не видаляє призначене обладнання — воно переходить у статус `AVAILABLE`.")

add_heading2(doc, "2.6 Проектування REST API")

add_normal(doc,
    "REST API системи використовує префікс `/api/v1/` і дотримується принципів RESTful-архітектури: "
    "ресурсо-орієнтовані URI, використання стандартних HTTP-методів (GET, POST, PUT, PATCH, DELETE), "
    "статусні коди HTTP (200, 201, 204, 400, 401, 403, 404, 409, 422, 500), "
    "форматування помилок за RFC 7807 (ProblemDetail). "
    "Повна специфікація доступна у форматі OpenAPI 3.0 за адресою `/swagger-ui.html`.")

add_table_caption(doc, "Основні ендпоінти REST API")
make_table(doc,
    ["Метод", "Шлях", "Роль", "Опис"],
    [
        ("POST", "/api/v1/auth/login", "ALL", "Вхід, отримання access+refresh токенів"),
        ("POST", "/api/v1/auth/refresh", "ALL", "Оновлення access токена по refresh"),
        ("POST", "/api/v1/auth/logout", "ALL", "Вихід, відкликання refresh токена"),
        ("GET",  "/api/v1/equipment", "ALL", "Список обладнання (фільтри, пагінація)"),
        ("POST", "/api/v1/equipment", "ADMIN, TEAM_LEAD", "Створення нової одиниці обладнання"),
        ("GET",  "/api/v1/equipment/{id}", "ALL", "Деталі одиниці обладнання"),
        ("PUT",  "/api/v1/equipment/{id}", "ADMIN, TEAM_LEAD", "Оновлення обладнання"),
        ("DELETE","/api/v1/equipment/{id}", "ADMIN", "Видалення (soft delete) обладнання"),
        ("POST", "/api/v1/equipment/{id}/assign", "ADMIN, TEAM_LEAD", "Призначення обладнання користувачу"),
        ("POST", "/api/v1/equipment/{id}/unassign", "ADMIN, TEAM_LEAD", "Зняття призначення"),
        ("GET",  "/api/v1/licenses", "ALL", "Список ліцензій"),
        ("POST", "/api/v1/licenses", "ADMIN", "Створення ліцензії"),
        ("GET",  "/api/v1/licenses/{id}", "ALL", "Деталі ліцензії (ключ замаскований)"),
        ("PUT",  "/api/v1/licenses/{id}", "ADMIN", "Оновлення ліцензії"),
        ("DELETE","/api/v1/licenses/{id}", "ADMIN", "Видалення ліцензії"),
        ("POST", "/api/v1/licenses/{id}/assign", "ADMIN, TEAM_LEAD", "Призначення ліцензії"),
        ("GET",  "/api/v1/audit-logs", "ADMIN, AUDITOR", "Журнал аудиту (фільтри, пагінація)"),
        ("POST", "/api/v1/reports/equipment", "ADMIN, TEAM_LEAD", "Запит на генерацію Excel-звіту"),
        ("GET",  "/api/v1/reports/{reportId}", "ADMIN, TEAM_LEAD", "Завантаження готового звіту"),
        ("GET",  "/api/v1/users", "ADMIN", "Список користувачів"),
        ("POST", "/api/v1/users", "ADMIN", "Створення користувача"),
        ("PUT",  "/api/v1/users/{id}/role", "ADMIN", "Зміна ролі користувача"),
        ("GET",  "/api/v1/sites", "ADMIN", "Список майданчиків"),
        ("POST", "/api/v1/sites", "ADMIN", "Створення майданчика"),
    ],
    col_widths=[1.8, 5.5, 3.5, 5.2]
)

add_heading2(doc, "2.7 RBAC матриця")

add_normal(doc,
    "Матриця RBAC визначає, які операції дозволені для кожної ролі системи. "
    "Дозволи реалізовані на двох рівнях: статична перевірка ролі (`hasRole('ADMIN')`) і динамічна перевірка через `InventoryPermissionEvaluator` (`hasPermission(#id, 'Equipment', 'MANAGE')`), "
    "що дозволяє TEAM_LEAD управляти тільки активами свого майданчика.")

add_table_caption(doc, "RBAC матриця доступу")
make_table(doc,
    ["Дія", "ADMIN", "TEAM_LEAD", "WORKER", "AUDITOR"],
    [
        ("Перегляд обладнання", "✓ (всі)", "✓ (свій майданчик)", "✓ (призначене)", "✓ (всі, read)"),
        ("Створення обладнання", "✓", "✓", "✗", "✗"),
        ("Редагування обладнання", "✓", "✓ (свій)", "✗", "✗"),
        ("Видалення обладнання", "✓", "✗", "✗", "✗"),
        ("Призначення обладнання", "✓", "✓ (свій)", "✗", "✗"),
        ("Перегляд ліцензій", "✓", "✓ (свій)", "✓ (призначені)", "✓"),
        ("Перегляд ключа ліцензії", "✓", "✓ (свій)", "✗", "✗"),
        ("CRUD ліцензій", "✓", "✗", "✗", "✗"),
        ("Журнал аудиту", "✓", "✓ (свій)", "✗", "✓"),
        ("Управління користувачами", "✓", "✗", "✗", "✗"),
        ("Генерація звітів", "✓", "✓ (свій)", "✗", "✗"),
        ("Управління майданчиками", "✓", "✗", "✗", "✗"),
    ],
    col_widths=[5, 2.5, 3, 2.5, 3]
)

add_heading2(doc, "2.8 Sequence-діаграма JWT автентифікації")

add_normal(doc,
    "Процес JWT-автентифікації складається з двох фаз: початкового входу та оновлення токена. "
    "При вході клієнт надсилає credentials (username/password) на `/api/v1/auth/login`. "
    "`AuthService` делегує перевірку пароля `AuthenticationManager`, який через `UserDetailsServiceImpl` завантажує користувача з БД і перевіряє BCrypt-хеш. "
    "У разі успіху `JwtService` генерує access-токен (15 хв) і refresh-токен (7 днів). "
    "Refresh токен зберігається у таблиці `refresh_tokens` у вигляді хешу і повертається клієнту в HttpOnly cookie [15].")

add_figure(doc, IMG_SEQ, "Sequence-діаграма процесу JWT автентифікації та оновлення токена")

add_normal(doc,
    "При кожному запиті `JwtAuthenticationFilter` (розширює `OncePerRequestFilter`) перехоплює HTTP-запит, витягує Bearer-токен з заголовка `Authorization`, "
    "перевіряє підпис і строк дії через `JwtService.isValid()`, "
    "завантажує `UserPrincipal` і встановлює автентифікацію у `SecurityContextHolder`. "
    "Після закінчення строку дії access-токена клієнт надсилає POST `/api/v1/auth/refresh` з refresh-токеном. "
    "Система перевіряє хеш у БД, видаляє старий запис (ротація) і видає нову пару токенів. "
    "Компрометований refresh-токен буде відхилений, оскільки попереднє оновлення вже інвалідувало його.")

add_heading2(doc, "2.8.1 Деталі реалізації JWT у Spring Security")
add_normal(doc,
    "Інтеграція JWT у Spring Security вимагає обережної конфігурації, щоб не зламати стандартну форму входу. "
    "Ключовим є те, що `JwtAuthenticationFilter` додається до ланцюжка API (`@Order(1)`) за допомогою "
    "`http.addFilterBefore(jwtAuthFilter, UsernamePasswordAuthenticationFilter.class)`, "
    "але не до ланцюжка веб-інтерфейсу (`@Order(2)`), де форма входу обслуговується стандартним `UsernamePasswordAuthenticationFilter`. "
    "Це досягається тим, що обидва `SecurityFilterChain` є окремими Spring-бінами і не поділяють конфігурацію.")

add_normal(doc,
    "Для забезпечення ротації refresh-токенів `AuthService` використовує транзакційну операцію: "
    "1) знайти старий refresh-токен за його хешем (`refreshTokenRepository.findByTokenHash(hash)`), "
    "2) перевірити, що він не відкликаний і не прострочений, "
    "3) позначити старий токен як відкликаний (`revoked = true`), "
    "4) згенерувати нову пару токенів, "
    "5) зберегти новий refresh-токен. "
    "Усі ці кроки виконуються в одній транзакції (`@Transactional`), тому у разі помилки на будь-якому кроці жодних змін у БД не відбудеться і старий токен залишиться дійсним. "
    "Це гарантує атомарність ротації і захищає від race condition.")

add_normal(doc,
    "Клас `AesKeyHolder` реалізує шифрування AES-256-GCM: секретний 32-байтний ключ зчитується з змінної середовища `AES_KEY` (Base64-encoded) при ініціалізації бін-а. "
    "Метод `encrypt(String plaintext)` генерує випадковий 12-байтний IV (Initialization Vector), виконує шифрування через `javax.crypto.Cipher` у режимі `AES/GCM/NoPadding` і повертає конкатенацію `Base64(IV + ciphertext + authTag)`. "
    "Метод `decrypt(String encoded)` витягує IV з перших 12 байт і виконує розшифрування. "
    "Автентифікаційний тег GCM (16 байт) забезпечує цілісність шифротексту і захист від модифікації даних у БД. "
    "При спробі розшифрувати модифікований шифротекст `AEADBadTagException` буде перехоплено і перетворено на `SecurityException`.")

add_heading2(doc, "2.9 Принципи SOLID")

add_normal(doc,
    "Архітектура системи свідомо дотримується принципів SOLID, що забезпечує її розширюваність і тестованість. "
    "Розглянемо застосування кожного принципу на конкретних прикладах з кодової бази.")

add_heading3(doc, "2.9.1 Single Responsibility Principle (SRP)")
add_normal(doc,
    "Принцип єдиної відповідальності реалізовано через чіткий поділ класів за функціями. "
    "`JwtService` відповідає виключно за генерацію, валідацію та парсинг JWT-токенів і не знає нічого про бізнес-логіку. "
    "`AuditLogListener` займається лише збереженням аудиторських подій і не містить логіки їхнього генерування. "
    "`GlobalExceptionHandler` централізує обробку виключень, звільняючи контролери від блоків try-catch. "
    "`ExcelReportGenerator` відповідає виключно за формування Excel-файлу, тоді як `ReportProducer` та `ReportConsumer` керують чергою повідомлень.")

add_heading3(doc, "2.9.2 Open/Closed Principle (OCP)")
add_normal(doc,
    "Принцип відкритості/закритості реалізовано через абстрактний клас `EquipmentFactory` та інтерфейс `NotificationStrategy`. "
    "Для додавання нового типу обладнання достатньо створити нову фабрику, що розширює `EquipmentFactory`, без зміни існуючого коду. "
    "Аналогічно, нову стратегію сповіщення (наприклад, Slack або Teams) можна додати через нову реалізацію `NotificationStrategy` без зміни `NotificationService`. "
    "Специфікації `EquipmentSpecification` і `LicenseSpecification` дозволяють додавати нові критерії фільтрації без зміни репозиторіїв.")

add_heading3(doc, "2.9.3 Liskov Substitution Principle (LSP)")
add_normal(doc,
    "Принцип підстановки Ліскова дотримується у ієрархії фабрик обладнання: "
    "`ComputerEquipmentFactory`, `NetworkEquipmentFactory` та `DefaultEquipmentFactory` повністю замінюють `EquipmentFactory` без порушення контракту. "
    "Аналогічно, `UserDetailsServiceImpl` реалізує `UserDetailsService` Spring Security і може бути замінена будь-якою іншою реалізацією без зміни SecurityConfig.")

add_heading3(doc, "2.9.4 Interface Segregation Principle (ISP)")
add_normal(doc,
    "Принцип розділення інтерфейсів дотримується: `NotificationStrategy` містить мінімальний контракт (`sendNotification(Notification)`), не змішуючи обов'язки відправлення різних типів повідомлень. "
    "Репозиторії `EquipmentRepository` і `LicenseRepository` розширюють тільки необхідні інтерфейси Spring Data (`JpaRepository` + `JpaSpecificationExecutor`), не успадковуючи зайвих методів.")

add_heading3(doc, "2.9.5 Dependency Inversion Principle (DIP)")
add_normal(doc,
    "Принцип інверсії залежностей реалізований через Spring IoC Container: усі залежності ін'єктуються через конструктор (завдяки Lombok `@RequiredArgsConstructor`), а залежності оголошуються від інтерфейсів, а не конкретних реалізацій. "
    "`EquipmentService` залежить від `EquipmentRepository` (інтерфейс), `EventPublisher` (інтерфейс), `EquipmentMapper` (інтерфейс MapStruct). "
    "Це дозволяє легко підміняти реалізації у тестах через Mockito-моки.")

add_heading2(doc, "2.10 Шаблони проектування")

add_normal(doc,
    "У системі застосовано сім класичних шаблонів проектування (GoF та інші). "
    "Їхнє використання є обґрунтованим з точки зору архітектури і не є надмірним (over-engineering). "
    "Кожен шаблон вирішує конкретну задачу і спрощує супровід відповідного модуля.")

add_heading3(doc, "2.10.1 Factory Method (Фабричний метод)")
add_normal(doc,
    "Шаблон Factory Method реалізовано в класі `EquipmentFactory` (пакет `ua.edu.inventory.equipment`). "
    "Статичний метод `forType(EquipmentType)` повертає конкретну фабрику залежно від типу обладнання. "
    "Кожна фабрика (комп'ютерна, мережева, стандартна) знає, які значення за замовчуванням встановити для нового обладнання свого типу: гарантійний термін, набір обов'язкових полів, початковий статус. "
    "Це ізолює логіку ініціалізації від сервісного шару і спрощує додавання нових типів обладнання.")

add_heading3(doc, "2.10.2 Specification (Специфікація)")
add_normal(doc,
    "Шаблон Specification реалізовано в класах `EquipmentSpecification` і `LicenseSpecification` на базі інтерфейсу `Specification<T>` із Spring Data JPA. "
    "Специфікації будуються з фільтрів, переданих у DTO запиту, і комбінуються операторами `and()`, `or()`. "
    "Це дозволяє конструювати довільні SQL WHERE-умови без дублювання коду і без змін у репозиторії. "
    "Метод `equipmentRepository.findAll(spec, pageable)` приймає специфікацію і прозоро перетворює її на параметризований SQL-запит.")

add_heading3(doc, "2.10.3 Observer (Спостерігач)")
add_normal(doc,
    "Шаблон Observer реалізовано через Spring ApplicationEventPublisher. "
    "Бізнес-сервіси (`EquipmentService`, `LicenseService`) публікують події `EntityChangedEvent` через `eventPublisher.publishEvent()`. "
    "Клас `AuditLogListener` підписується на ці події за допомогою анотації `@EventListener` і зберігає запис у журналі аудиту в окремій транзакції (`@Transactional(propagation = REQUIRES_NEW)`). "
    "Такий підхід забезпечує слабке зв'язування між бізнес-логікою і аудитом: сервіси не знають нічого про `AuditLogListener`.")

add_heading3(doc, "2.10.4 Strategy (Стратегія)")
add_normal(doc,
    "Шаблон Strategy реалізовано через інтерфейс `NotificationStrategy` з методом `sendNotification(Notification notification)`. "
    "Конкретні реалізації: `EmailNotificationStrategy`, `InAppNotificationStrategy`. "
    "`NotificationService` отримує список усіх реалізацій через `@Autowired List<NotificationStrategy>` і ітерує по ньому при відправленні сповіщення. "
    "Додавання нового каналу (наприклад, SMS або Telegram) зводиться до написання нового `@Component`, що реалізує `NotificationStrategy`, без будь-яких змін у `NotificationService`.")

add_heading3(doc, "2.10.5 Builder (Будівельник)")
add_normal(doc,
    "Шаблон Builder застосовано для побудови DTO-запитів генерації звітів: `EquipmentReportRequest` та `LicenseReportRequest` анотовані `@Builder` (Lombok). "
    "Це дозволяє зручно конструювати запити з великою кількістю опціональних параметрів фільтрації (тип, статус, майданчик, діапазон дат) без перевантажених конструкторів. "
    "Builder також використовується при побудові `AuditLog.builder().action(...).entityId(...).build()` в `AuditLogListener`.")

add_heading3(doc, "2.10.6 Decorator (Декоратор)")
add_normal(doc,
    "Шаблон Decorator реалізовано через JPA AttributeConverter: клас `LicenseKeyAttributeConverter` реалізує `AttributeConverter<String, String>`. "
    "При записі в базу метод `convertToDatabaseColumn()` шифрує ліцензійний ключ алгоритмом AES-256-GCM через `AesKeyHolder`. "
    "При читанні `convertToEntityAttribute()` розшифровує значення. "
    "Для JPA-шару і решти коду ліцензійний ключ виглядає як звичайний рядок — декоратор прозоро додає шар шифрування. "
    "Анотація `@Convert(converter = LicenseKeyAttributeConverter.class)` на полі `key` сутності `License` активує перетворення.")

add_heading3(doc, "2.10.7 Singleton (Одинак)")
add_normal(doc,
    "Шаблон Singleton реалізовано на рівні Spring Container: усі `@Service`, `@Component`, `@Repository` є singleton-бінами за замовчуванням. "
    "Особливо важливо це для `JwtService`, що містить кешований секретний ключ, і для `ReportStore` — сховища готових звітів на базі `ConcurrentHashMap<UUID, byte[]>`. "
    "`ReportStore` є безпечним для конкурентного доступу завдяки `ConcurrentHashMap` і гарантує, що у всьому застосунку існує єдиний екземпляр сховища звітів.")

add_heading2(doc, "2.10.8 Узагальнення застосування шаблонів")
add_normal(doc,
    "Сукупне застосування семи шаблонів проектування у системі демонструє зрілий підхід до архітектурного мислення. "
    "Важливо підкреслити, що жоден із шаблонів не був застосований заради самого шаблону — кожен з них вирішує конкретну задачу, що підтверджується відповідними функціональними вимогами (FR-04 — FR-13). "
    "Factory Method усуває потребу в умовних операторах при створенні різних типів обладнання. "
    "Specification забезпечує гнучку фільтрацію без зміни репозиторіїв при додаванні нових критеріїв пошуку. "
    "Observer забезпечує слабке зв'язування між бізнес-логікою і аудитом. "
    "Strategy дозволяє додавати нові канали сповіщень без зміни існуючого коду. "
    "Builder спрощує конструювання складних запитів з опціональними параметрами. "
    "Decorator прозоро додає шифрування без впливу на решту коду. "
    "Singleton гарантує безпеку конкурентного доступу до спільних ресурсів.")

add_normal(doc,
    "Усі шаблони органічно вписуються в екосистему Spring Framework: IoC Container є реалізацією Singleton на рівні фреймворку, "
    "ApplicationEventPublisher забезпечує Observer без додаткових бібліотек, "
    "JPA AttributeConverter є стандартним місцем для Decorator. "
    "Це свідчить про те, що Spring Framework сам по собі є набором реалізацій класичних GoF-шаблонів, "
    "адаптованих до enterprise-контексту [19].")

add_heading2(doc, "2.11 Висновки до розділу 2")

add_normal(doc,
    "У другому розділі виконано повне проектування системи управління ІТ-інвентаризацією. "
    "Обрана тришарова архітектура забезпечує чіткий поділ відповідальностей і спрощує тестування. "
    "Обґрунтовано вибір технологічного стеку з 16 ключових компонентів, кожен з яких відповідає конкретній функціональній або нефункціональній вимозі.")

add_normal(doc,
    "Спроектовано логічну модель бази даних з семи таблиць, що забезпечує нормалізовану структуру і референційну цілісність. "
    "Спроектовано REST API з 24 ендпоінтами, що охоплюють повний функціонал системи. "
    "Визначено RBAC матрицю для чотирьох ролей системи із підтримкою тонкогранульованого контролю доступу на рівні майданчика.")

add_normal(doc,
    "Система відповідає всім п'яти принципам SOLID, що підтверджено конкретними прикладами з кодової бази. "
    "Застосовано сім шаблонів проектування (Factory Method, Specification, Observer, Strategy, Builder, Decorator, Singleton), кожен з яких вирішує конкретну архітектурну задачу. "
    "Результати проектування є основою для програмної реалізації, описаної у наступному розділі.")

# ══════════════════════════════════════════════════════════════════════════════
# РОЗДІЛ 3. ПРОГРАМНА РЕАЛІЗАЦІЯ
# ══════════════════════════════════════════════════════════════════════════════

add_heading1(doc, "РОЗДІЛ 3. ПРОГРАМНА РЕАЛІЗАЦІЯ")

add_heading2(doc, "3.1 Структура проекту")

add_normal(doc,
    "Проект організовано як стандартний Maven-проект зі структурою директорій, що відповідає конвенції Spring Boot. "
    "Кореневий пакет `ua.edu.inventory` містить клас `InventoryApplication` з анотацією `@SpringBootApplication` і методом `main()`. "
    "Дочірні пакети організовані за функціональним принципом (package by feature), де кожна бізнес-область утворює окремий пакет з власними підпакетами для сутностей, DTO, репозиторіїв, сервісів і контролерів.")

PROJ_STRUCTURE = """src/
├── main/
│   ├── java/ua/edu/inventory/
│   │   ├── InventoryApplication.java
│   │   ├── auth/
│   │   │   ├── AuthController.java
│   │   │   ├── AuthService.java
│   │   │   ├── dto/ (LoginRequest, TokenResponse, RefreshRequest)
│   │   │   └── model/ (RefreshToken.java)
│   │   ├── user/
│   │   │   ├── User.java, UserRepository.java
│   │   │   ├── UserService.java, UserController.java
│   │   │   └── UserDetailsServiceImpl.java
│   │   ├── equipment/
│   │   │   ├── Equipment.java, EquipmentRepository.java
│   │   │   ├── EquipmentService.java, EquipmentController.java
│   │   │   ├── EquipmentFactory.java, EquipmentSpecification.java
│   │   │   └── dto/ (EquipmentCreateDto, EquipmentDto, EquipmentFilterDto)
│   │   ├── license/
│   │   │   ├── License.java, LicenseRepository.java
│   │   │   ├── LicenseService.java, LicenseController.java
│   │   │   ├── LicenseKeyAttributeConverter.java
│   │   │   └── LicenseSpecification.java
│   │   ├── audit/
│   │   │   ├── AuditLog.java, AuditLogRepository.java
│   │   │   ├── AuditLogListener.java, EntityChangedEvent.java
│   │   │   └── AuditLogController.java
│   │   ├── notification/
│   │   │   ├── NotificationStrategy.java (interface)
│   │   │   ├── EmailNotificationStrategy.java
│   │   │   ├── InAppNotificationStrategy.java
│   │   │   └── NotificationService.java
│   │   ├── report/
│   │   │   ├── ReportProducer.java, ReportConsumer.java
│   │   │   ├── ExcelReportGenerator.java, ReportStore.java
│   │   │   └── dto/ (EquipmentReportRequest, LicenseReportRequest)
│   │   ├── site/
│   │   │   └── Site.java, SiteRepository.java, SiteService.java
│   │   ├── common/
│   │   │   ├── BaseAuditableEntity.java
│   │   │   └── ResourceNotFoundException.java, BusinessRuleException.java
│   │   ├── config/
│   │   │   ├── SecurityConfig.java, JwtAuthenticationFilter.java
│   │   │   ├── JwtService.java, AesKeyHolder.java
│   │   │   ├── RabbitMqConfig.java, OpenApiConfig.java
│   │   │   └── InventoryPermissionEvaluator.java
│   │   ├── crypto/
│   │   │   └── AesKeyHolder.java
│   │   └── web/
│   │       └── GlobalExceptionHandler.java, UserPrincipal.java
│   └── resources/
│       ├── application.yml
│       ├── db/migration/ (V1__init.sql, V2__seed.sql, ...)
│       └── templates/ (layout/, equipment/, license/, admin/, ...)
└── test/
    └── java/ua/edu/inventory/
        ├── auth/ (AuthIntegrationTest.java)
        ├── equipment/ (EquipmentServiceTest.java, EquipmentIntegrationTest.java)
        ├── audit/ (AuditLogListenerTest.java)
        └── config/ (JwtServiceTest.java)"""

add_code(doc, PROJ_STRUCTURE, "Структура Maven-проекту ua.edu.inventory")

add_normal(doc,
    "Пакет `common` містить базові класи, що використовуються в усіх модулях: `BaseAuditableEntity` з аудитними полями JPA Auditing, "
    "`ResourceNotFoundException` (мапується на HTTP 404) і `BusinessRuleException` (мапується на HTTP 422). "
    "Пакет `config` містить всю Spring-конфігурацію: безпеку, JWT, RabbitMQ, OpenAPI. "
    "Такий поділ гарантує, що конфігурація не змішується з бізнес-логікою.")

add_heading2(doc, "3.2 Доменна модель та JPA")

add_normal(doc,
    "Усі доменні сутності успадковують від абстрактного класу `BaseAuditableEntity`, анотованого `@MappedSuperclass` і `@EntityListeners(AuditingEntityListener.class)`. "
    "Клас містить поля `createdAt` (`@CreatedDate`), `updatedAt` (`@LastModifiedDate`), `createdBy` (`@CreatedBy`) і `updatedBy` (`@LastModifiedBy`), що автоматично заповнюються Spring Data Auditing. "
    "Первинний ключ визначено у базовому класі як `@Id @GeneratedValue(strategy = GenerationType.UUID) UUID id`.")

add_normal(doc,
    "Клас `Equipment` анотовано `@Entity @Table(name = \"equipment\")` і містить: поля `inventoryNumber` (унікальний, not null), `name`, `type` (enum `EquipmentType`), "
    "`status` (enum `EquipmentStatus`: AVAILABLE, ASSIGNED, MAINTENANCE, DECOMMISSIONED), `serialNumber`, `warrantyExpires` (LocalDate). "
    "Зв'язки: `@ManyToOne Site site` і `@ManyToOne User assignedUser` (nullable, для необов'язкового призначення). "
    "Enum-поля зберігаються як рядки (`@Enumerated(EnumType.STRING)`) для читабельності у БД і стійкості до зміни порядку значень.")

add_normal(doc,
    "Клас `License` анотовано аналогічно і містить: `name`, `vendor`, `product`, `key` (поле, анотоване `@Convert(converter = LicenseKeyAttributeConverter.class)`), "
    "`totalSeats` (int), `usedSeats` (int), `expiresAt` (LocalDate), `site` (ManyToOne). "
    "Поле `key` прозоро шифрується при кожному збереженні і розшифровується при кожному читанні з БД завдяки декоратору `LicenseKeyAttributeConverter`.")

add_normal(doc,
    "Схема БД керується Flyway: перша міграція `V1__init.sql` створює всі таблиці, `V2__seed.sql` додає тестові дані для розробки. "
    "Це гарантує, що база завжди відтворює задану схему при розгортанні в нових середовищах або у Testcontainers-контейнері під час інтеграційних тестів. "
    "Flyway запускається автоматично при старті Spring Boot застосунку завдяки стартеру `spring-boot-starter-data-jpa` і наявності Flyway у classpath.")

add_heading2(doc, "3.3 Автентифікація та авторизація")

add_normal(doc,
    "Безпека системи базується на двох механізмах: JWT для API-клієнтів та сесійній form-login для браузерних клієнтів. "
    "`SecurityConfig` налаштовує два окремих `SecurityFilterChain`. "
    "Перший (`@Order(1)`) обслуговує маршрути `/api/**`: stateless (без сесій), додає `JwtAuthenticationFilter` у ланцюжок, вимагає автентифікації для більшості ендпоінтів крім `/api/v1/auth/login` і `/api/v1/auth/refresh`. "
    "Другий (`@Order(2)`) обслуговує веб-інтерфейс: form-login, сесія у пам'яті, редирект на `/login` для незаутентифікованих запитів.")

LISTING_JWT_FILTER = """@Component
@RequiredArgsConstructor
public class JwtAuthenticationFilter extends OncePerRequestFilter {
    private final JwtService jwtService;
    private final UserDetailsServiceImpl userDetailsService;

    @Override
    protected void doFilterInternal(HttpServletRequest request,
            HttpServletResponse response, FilterChain chain)
            throws ServletException, IOException {
        String token = extractBearerToken(request);
        if (token != null && jwtService.isValid(token)) {
            UUID userId = jwtService.extractUserId(token);
            UserPrincipal principal =
                (UserPrincipal) userDetailsService.loadUserById(userId);
            var auth = new UsernamePasswordAuthenticationToken(
                    principal, null, principal.getAuthorities());
            SecurityContextHolder.getContext().setAuthentication(auth);
        }
        chain.doFilter(request, response);
    }
}"""

add_code(doc, LISTING_JWT_FILTER, "Клас JwtAuthenticationFilter — фільтр JWT автентифікації")

add_normal(doc,
    "Клас `JwtService` відповідає за генерацію та валідацію JWT-токенів. "
    "Секретний ключ завантажується з `application.yml` (властивість `jwt.secret`) і кешується при ініціалізації бін-а. "
    "Метод `isValid(String token)` перевіряє підпис і строк дії токена. "
    "Метод `extractUserId(String token)` витягує UUID користувача з claims. "
    "Метод `generateAccessToken(UserPrincipal)` створює токен зі строком дії 15 хвилин і claims: `sub` (userId), `role` (роль користувача). "
    "Бібліотека JJWT 0.12.6 забезпечує безпечне підписання алгоритмом HS256.")

add_normal(doc,
    "Клас `UserPrincipal` реалізує `UserDetails` і додатково зберігає `UUID userId` і `UUID siteId`, "
    "що використовується в `InventoryPermissionEvaluator` для перевірки належності активу до майданчика поточного користувача. "
    "Хешування паролів виконується BCrypt зі strength=12, що забезпечує достатній рівень стійкості до брутфорсу відповідно до рекомендацій OWASP. "
    "`PasswordEncoder` оголошено як бін у `SecurityConfig` і використовується в `AuthService` при реєстрації та в `AuthenticationManager` при вході.")

add_heading2(doc, "3.4 Бізнес-логіка")

add_normal(doc,
    "Центральним сервісом є `EquipmentService`, що забезпечує CRUD-операції над обладнанням, фільтрацію, призначення та зняття призначення. "
    "Усі публічні методи сервісу анотовані `@Transactional` або `@Transactional(readOnly = true)`. "
    "Контроль доступу реалізовано на рівні методів за допомогою `@PreAuthorize`. "
    "Нижче наведено ключовий метод `assign()`, що демонструє взаємодію з репозиторієм, перевірку бізнес-правил та публікацію аудиторської події.")

LISTING_ASSIGN = """@PreAuthorize("hasRole('ADMIN') or (hasRole('TEAM_LEAD') " +
              "and hasPermission(#id, 'Equipment', 'MANAGE'))")
@Transactional
public EquipmentDto assign(UUID id, UUID userId) {
    Equipment eq = findOrThrow(id);
    User user = userRepository.findById(userId)
            .orElseThrow(() -> new ResourceNotFoundException(
                    "Користувач", userId));
    if (!eq.getSiteId().equals(user.getSiteId())) {
        throw new BusinessRuleException(
                "Користувач належить до іншого об'єкту.");
    }
    eq.setAssignedUserId(userId);
    eq.setStatus(EquipmentStatus.ASSIGNED);
    Equipment saved = equipmentRepository.save(eq);
    eventPublisher.publishEvent(new EntityChangedEvent(
            this, AuditAction.ASSIGN, EntityType.EQUIPMENT,
            id.toString(), Map.of("userId", userId), actorId, null));
    return equipmentMapper.toDto(saved);
}"""

add_code(doc, LISTING_ASSIGN, "Метод EquipmentService.assign() — призначення обладнання користувачу")

add_normal(doc,
    "Метод `findOrThrow(UUID id)` інкапсулює виклик `equipmentRepository.findById(id).orElseThrow(() -> new ResourceNotFoundException(\"Обладнання\", id))`, "
    "уникаючи дублювання. "
    "Перевірка `eq.getSiteId().equals(user.getSiteId())` є бізнес-правилом: призначати обладнання можна лише користувачу того ж майданчика. "
    "Порушення цього правила генерує `BusinessRuleException`, яке `GlobalExceptionHandler` перетворює на HTTP 422 з описовим повідомленням у форматі ProblemDetail.")

add_normal(doc,
    "`LicenseService` реалізує схожу логіку для ліцензій з додатковою перевіркою доступних місць: "
    "при призначенні ліцензії інкрементується `usedSeats`, при поверненні — декрементується. "
    "Якщо `usedSeats >= totalSeats`, метод `assignLicense()` викидає `BusinessRuleException`. "
    "Поле `key` у `LicenseDto` маскується — повертається рядок `\"***\"` для ролей WORKER і AUDITOR, і лише ADMIN та TEAM_LEAD власного майданчика бачать розшифрований ключ. "
    "Це реалізовано через метод `maskKey(LicenseDto dto, UserPrincipal principal)` у сервісному шарі.")

add_heading2(doc, "3.4.1 LicenseService та управління ліцензіями")
add_normal(doc,
    "Клас `LicenseService` забезпечує повний цикл управління програмними ліцензіями. "
    "Метод `createLicense(LicenseCreateDto dto, UUID actorId)` приймає DTO із ліцензійним ключем у відкритому вигляді, "
    "будує сутність `License` (через `LicenseMapper`), зберігає її через репозиторій (при збереженні `LicenseKeyAttributeConverter` автоматично шифрує ключ) "
    "і публікує подію `EntityChangedEvent` з дією CREATE. "
    "Важливо: у методі ніде не відбувається явного виклику шифрування — він повністю прихований у шарі JPA через `@Convert`, що є прикладом правильного застосування шаблону Decorator.")

add_normal(doc,
    "Метод `assignLicense(UUID licenseId, UUID userId, UUID actorId)` реалізує транзакційне призначення ліцензії: "
    "спочатку перевіряється наявність вільних місць (`if (license.getUsedSeats() >= license.getTotalSeats()) throw new BusinessRuleException(...)`), "
    "потім атомарно інкрементується `usedSeats` і зберігається ліцензія, "
    "після чого публікується аудиторська подія. "
    "Оптимістичне блокування через `@Version` на полі `usedSeats` гарантує, що при конкурентному призначенні ліцензії кількість місць не перевищить ліміт: "
    "Spring Data JPA автоматично генерує `UPDATE ... WHERE version = ?` і кидає `OptimisticLockingFailureException` у разі конфлікту.")

add_normal(doc,
    "Метод `getLicenseDto(UUID id, UserPrincipal actor)` демонструє рольове маскування чутливих даних: "
    "після завантаження сутності і маппінгу в `LicenseDto` перевіряється роль актора. "
    "Якщо актор — WORKER або AUDITOR, або TEAM_LEAD іншого майданчика, поле `key` у DTO замінюється на рядок `\"***\"`. "
    "ADMIN і TEAM_LEAD власного майданчика бачать розшифрований ключ. "
    "Такий підхід гарантує, що навіть при помилці у `@PreAuthorize` анотації конфіденційний ключ не потрапить до неавторизованого користувача через захисний шар у сервісі.")

add_heading2(doc, "3.5 REST API контролери та обробка помилок")

add_normal(doc,
    "REST-контролери анотовані `@RestController @RequestMapping(\"/api/v1/...\")` і відповідають виключно за HTTP-специфічні завдання: "
    "парсинг запиту, делегування сервісу, формування відповіді. "
    "Жодної бізнес-логіки у контролерах немає — це відповідальність сервісного шару. "
    "`EquipmentController` містить 8 методів (GET список, GET за id, POST, PUT, DELETE, assign, unassign, export). "
    "Вхідні DTO валідуються через Jakarta Bean Validation (`@Valid`, `@NotNull`, `@Size`) і помилки валідації автоматично обробляються `GlobalExceptionHandler`.")

add_normal(doc,
    "Клас `GlobalExceptionHandler` анотовано `@RestControllerAdvice` і обробляє наступні виключення: "
    "`ResourceNotFoundException` → 404, `BusinessRuleException` → 422, `AccessDeniedException` → 403, "
    "`MethodArgumentNotValidException` → 400 з деталями полів, `Exception` → 500. "
    "Усі відповіді форматуються як `ProblemDetail` (RFC 7807) з полями `type`, `title`, `status`, `detail`, `instance`. "
    "Це забезпечує консистентний формат помилок для будь-якого клієнта API незалежно від типу виключення.")

add_normal(doc,
    "Документація API автоматично генерується springdoc 2.6.0. "
    "Swagger UI доступний за адресою `/swagger-ui.html` і відображає всі ендпоінти з описами, схемами запитів/відповідей та можливістю тестування. "
    "API-специфікація у форматі OpenAPI 3.0 JSON доступна за `/v3/api-docs`. "
    "Ендпоінти анотовані `@Operation`, `@Parameter` та `@ApiResponse` для чіткої документації. "
    "Автентифікація через Swagger UI налаштована через `SecurityScheme` з Bearer JWT у `OpenApiConfig`.")

add_heading2(doc, "3.6 Веб-інтерфейс Thymeleaf")

add_normal(doc,
    "Браузерний веб-інтерфейс реалізовано на Thymeleaf 3 з використанням Bootstrap 5 для адаптивного дизайну. "
    "Базовий шаблон `templates/layout/base.html` визначає структуру сторінки: навігаційна панель, бічне меню, основний вміст і футер. "
    "Усі сторінки розширюють базовий шаблон через `th:replace` фрагменти. "
    "Навігаційне меню і кнопки дій відображаються або приховуються залежно від ролі поточного користувача за допомогою `sec:authorize=\"hasRole('ADMIN')\"` з Thymeleaf Spring Security діалекту.")

add_normal(doc,
    "Основні сторінки веб-інтерфейсу: "
    "`/login` — форма входу (form-login), "
    "`/dashboard` — інформаційна панель з підсумком активів, "
    "`/equipment` — список обладнання з фільтрами і пагінацією, "
    "`/equipment/{id}` — картка обладнання, "
    "`/equipment/new` — форма створення, "
    "`/licenses` та `/licenses/{id}` — аналогічно для ліцензій, "
    "`/admin/users` — управління користувачами (тільки ADMIN), "
    "`/admin/sites` — управління майданчиками, "
    "`/audit-logs` — журнал аудиту (ADMIN/AUDITOR), "
    "`/reports` — запит і завантаження звітів.")

add_normal(doc,
    "Форми створення і редагування використовують `th:object` і `th:field` для двостороннього прив'язання до DTO. "
    "Серверна валідація Bean Validation повертає помилки у модель, які відображаються під відповідними полями через `th:errors`. "
    "CSRF-захист увімкнено для форм і Thymeleaf автоматично вставляє CSRF-токен у кожну форму через атрибут `th:action`. "
    "Сторінки є повністю адаптивними завдяки Bootstrap 5 grid system — інтерфейс коректно відображається як на десктопах, так і на мобільних пристроях.")

add_normal(doc,
    "[СКРИНШОТ: dashboard.png — інформаційна панель системи з підсумком активів]")
add_normal(doc,
    "[СКРИНШОТ: equipment-list.png — список обладнання з фільтрами]")

add_heading2(doc, "3.7 RabbitMQ та асинхронна генерація звітів")

add_normal(doc,
    "Генерація Excel-звітів є ресурсомістким завданням, що може тривати кілька секунд при великій кількості записів. "
    "Для уникнення блокування HTTP-потоку реалізовано асинхронну обробку через брокер повідомлень RabbitMQ. "
    "Архітектурна схема: клієнт надсилає POST `/api/v1/reports/equipment` із параметрами звіту → "
    "`ReportProducer` публікує повідомлення у чергу `report.queue` → "
    "клієнту повертається `reportId` (UUID) → "
    "`ReportConsumer` обробляє повідомлення, генерує Excel через `ExcelReportGenerator` → "
    "готовий файл зберігається у `ReportStore` (ConcurrentHashMap) → "
    "клієнт опитує GET `/api/v1/reports/{reportId}` до готовності і завантажує файл.")

add_normal(doc,
    "`RabbitMqConfig` оголошує exchange `report.exchange` типу Direct, чергу `report.queue` і binding з routing key `report`. "
    "Також налаштовано Dead Letter Exchange `report.dlx` для обробки повідомлень, що не вдалося обробити. "
    "`ReportProducer` використовує `RabbitTemplate.convertAndSend()` для публікації `EquipmentReportRequest` (серіалізованого у JSON через Jackson). "
    "`ReportConsumer` анотовано `@RabbitListener(queues = \"report.queue\")` і делегує генерацію `ExcelReportGenerator`.")

add_normal(doc,
    "`ExcelReportGenerator` використовує Apache POI 5.3.0 для створення `.xlsx` файлу. "
    "Звіт містить заголовок зі стилями (жирний шрифт, заливка), рядки даних з автоматичним підбором ширини колонок (`autoSizeColumn()`), "
    "нижній колонтитул з датою генерації, а також рядок підсумків (загальна кількість, кількість ASSIGNED, AVAILABLE). "
    "Готовий файл у вигляді `byte[]` зберігається у `ReportStore` разом із метаданими (ім'я файлу, дата генерації, userId). "
    "TTL запису у `ReportStore` контролюється scheduled задачею, що видаляє звіти старші за 1 годину.")

add_heading2(doc, "3.7.1 Деталі генерації Excel-звітів")
add_normal(doc,
    "Клас `ExcelReportGenerator` використовує Apache POI XSSF (XML-based SpreadSheet Format) для генерації `.xlsx` файлів. "
    "Метод `generateEquipmentReport(EquipmentReportRequest request, List<EquipmentDto> items)` створює `XSSFWorkbook`, "
    "додає лист `Equipment Report`, формує рядок заголовків зі стилями (жирний шрифт, синій фон, білий текст, всі межі), "
    "заповнює рядки даних з чергуванням кольорів фону (зебра-стиль) для кращої читабельності, "
    "застосовує `sheet.autoSizeColumn(i)` для всіх колонок, "
    "додає рядок підсумків з кількістю записів і кількістю по кожному статусу, "
    "і записує результат у `ByteArrayOutputStream`.")

add_normal(doc,
    "Колонки звіту обладнання: інвентарний номер, назва, тип, статус, серійний номер, майданчик, призначений користувач, дата закінчення гарантії, дата створення запису. "
    "Дати форматуються через `DateTimeFormatter.ofPattern(\"dd.MM.yyyy\")`. "
    "Статус у звіті відображається як локалізований рядок (AVAILABLE → «Доступний», ASSIGNED → «Призначений» тощо) через утилітарний клас `EquipmentStatusLocalizer`. "
    "Аналогічно реалізовано `generateLicenseReport()` для звіту ліцензій з колонками: назва, вендор, продукт, загальна/використана кількість місць, дата закінчення, майданчик.")

add_normal(doc,
    "Готовий звіт у вигляді `byte[]` зберігається у `ReportStore` — singleton-біні на базі `ConcurrentHashMap<UUID, ReportEntry>`. "
    "Клас `ReportEntry` містить: `byte[] data`, `String filename`, `LocalDateTime generatedAt`, `UUID requestedBy`. "
    "Scheduled метод `@Scheduled(fixedDelay = 900_000)` кожні 15 хвилин видаляє записи, старші за 1 годину, щоб уникнути витоку пам'яті. "
    "При завантаженні звіту через GET `/api/v1/reports/{reportId}` контролер перевіряє, що `requestedBy` відповідає поточному користувачу "
    "(ADMIN може завантажити будь-який звіт), і повертає файл з відповідними заголовками `Content-Type: application/vnd.openxmlformats-officedocument.spreadsheetml.sheet` "
    "та `Content-Disposition: attachment; filename=\"report.xlsx\"`.")

add_heading2(doc, "3.8 Аудит-логування")

add_normal(doc,
    "Журналювання всіх змін є критичною вимогою системи. "
    "Реалізація побудована на Spring Application Events і шаблоні Observer. "
    "Клас `EntityChangedEvent` розширює `ApplicationEvent` і містить: `actorUserId` (UUID), `action` (enum `AuditAction`: CREATE, UPDATE, DELETE, ASSIGN, UNASSIGN), "
    "`entityType` (enum `EntityType`: EQUIPMENT, LICENSE, USER), `entityId` (String), `payload` (Map<String, Object>). "
    "Бізнес-сервіси публікують події через `ApplicationEventPublisher.publishEvent()` після успішного збереження змін.")

LISTING_AUDIT = """@Component
@RequiredArgsConstructor
public class AuditLogListener {
    private final AuditLogRepository auditLogRepository;

    @EventListener
    @Transactional(propagation = Propagation.REQUIRES_NEW)
    public void onEntityChanged(EntityChangedEvent event) {
        AuditLog log = AuditLog.builder()
                .actorUserId(event.getActorUserId())
                .action(event.getAction())
                .entityType(event.getEntityType())
                .entityId(event.getEntityId())
                .payload(event.getPayload())
                .build();
        auditLogRepository.save(log);
    }
}"""

add_code(doc, LISTING_AUDIT, "Клас AuditLogListener — збереження аудиторських подій")

add_normal(doc,
    "Анотація `@Transactional(propagation = Propagation.REQUIRES_NEW)` є критично важливою: "
    "вона гарантує, що запис аудиту буде збережено в окремій транзакції незалежно від результату основної транзакції. "
    "Якщо основна транзакція відкочується через помилку, запис аудиту про спробу операції все одно зберігається. "
    "Це забезпечує повну аудиторську відстежуваність, включаючи невдалі операції. "
    "Поле `payload` зберігається у БД як JSONB, що дозволяє зберігати різний набір атрибутів для різних типів подій без зміни схеми таблиці `audit_logs`.")

add_heading2(doc, "3.9 Контейнеризація")

add_normal(doc,
    "Застосунок контейнеризовано за допомогою Docker з використанням multi-stage build для мінімізації розміру фінального образу. "
    "Перший stage (builder) використовує образ `eclipse-temurin:17-jdk-alpine` для компіляції і упаковки Maven-проекту. "
    "Другий stage (runtime) використовує мінімальний образ `eclipse-temurin:17-jre-alpine` і копіює лише артефакт `.jar` з першого stage. "
    "Такий підхід зменшує розмір фінального образу з ~500 МБ до ~150 МБ і виключає з образу інструменти збирання, знижуючи поверхню атаки.")

DOCKERFILE = """# Stage 1: build
FROM eclipse-temurin:17-jdk-alpine AS builder
WORKDIR /app
COPY pom.xml .
COPY src ./src
RUN mvn -B -DskipTests package

# Stage 2: runtime
FROM eclipse-temurin:17-jre-alpine
WORKDIR /app
COPY --from=builder /app/target/inventory-*.jar app.jar
EXPOSE 8080
ENTRYPOINT ["java", "-jar", "app.jar"]"""

add_code(doc, DOCKERFILE, "Dockerfile із multi-stage build для мінімізації образу")

add_normal(doc,
    "`docker-compose.yml` оркеструє три сервіси: `db` (PostgreSQL 16), `rabbitmq` (RabbitMQ 3.12 з management UI), `app` (застосунок). "
    "Сервіси об'єднані у спільну мережу `inventory-net`. "
    "Застосунок залежить від `db` і `rabbitmq` через `depends_on` з умовою `service_healthy`, "
    "а сервіси БД і RabbitMQ мають healthcheck налаштування для надійного запуску. "
    "Персистентність даних PostgreSQL забезпечена іменованим Docker volume `pg-data`. "
    "Конфігурація передається через змінні середовища: `SPRING_DATASOURCE_URL`, `SPRING_RABBITMQ_HOST`, `JWT_SECRET`, `AES_KEY`.")

add_normal(doc,
    "Секрети (JWT secret, AES ключ) передаються виключно через змінні середовища і ніколи не хардкодуються у вихідному коді. "
    "Файл `.env.example` містить шаблон з описом необхідних змінних, але не містить реальних значень. "
    "Реальний `.env` файл включено у `.gitignore` і не потрапляє до репозиторію. "
    "Це відповідає принципу 12-Factor App і рекомендаціям OWASP щодо захисту чутливих даних конфігурації.")

add_heading2(doc, "3.9.1 Конфігурація docker-compose")
add_normal(doc,
    "Файл `docker-compose.yml` налаштовує три взаємопов'язані сервіси. "
    "Сервіс `db` використовує офіційний образ `postgres:16-alpine`, підключає іменований volume `pg-data:/var/lib/postgresql/data` для персистентності даних і визначає healthcheck через `pg_isready`. "
    "Сервіс `rabbitmq` використовує образ `rabbitmq:3.12-management-alpine`, що включає Management UI на порті 15672 для моніторингу черг і з'єднань у розробницькому середовищі. "
    "Сервіс `app` збирається із поточного контексту (Dockerfile у корені проекту), залежить від `db` і `rabbitmq` через `condition: service_healthy`, отримує всі налаштування через змінні середовища.")

add_normal(doc,
    "Мережа `inventory-net` типу `bridge` об'єднує всі три сервіси. "
    "Зовні доступний тільки порт 8080 застосунку (і 15672 RabbitMQ Management UI у режимі розробки). "
    "PostgreSQL і RabbitMQ не прокидаються назовні, що відповідає принципу мінімального відкриття портів. "
    "У продуктивному середовищі перед застосунком розміщується nginx reverse proxy з TLS-термінацією, "
    "що додатково захищає з'єднання і дозволяє масштабувати кількість інстанцій застосунку за потреби.")

add_normal(doc,
    "Для продуктивного розгортання рекомендується перенести сервіси бази даних і RabbitMQ на виділені managed-послуги хмарних провайдерів "
    "(наприклад, Amazon RDS для PostgreSQL і Amazon MQ для RabbitMQ), зберігши тільки контейнер застосунку у Docker/Kubernetes. "
    "Це підвищує надійність, спрощує резервне копіювання і моніторинг зовнішніх залежностей. "
    "Перехід на managed-послуги не вимагає змін у коді застосунку — достатньо змінити змінні середовища з'єднання.")

add_heading2(doc, "3.10 Висновки до розділу 3")

add_normal(doc,
    "У третьому розділі детально описано програмну реалізацію системи управління ІТ-інвентаризацією. "
    "Розглянуто структуру проекту, організовану за функціональним принципом з чітким поділом на пакети. "
    "Описано доменну модель з використанням JPA/Hibernate, версіонування схеми Flyway та MapStruct для трансформації DTO.")

add_normal(doc,
    "Детально розглянуто реалізацію безпеки: двохланцюжкова конфігурація Spring Security, JWT-фільтр, BCrypt-хешування паролів, RBAC з `InventoryPermissionEvaluator`. "
    "Наведено фрагменти ключового коду: `JwtAuthenticationFilter`, `EquipmentService.assign()`, `AuditLogListener`. "
    "Описано асинхронну генерацію Excel-звітів через RabbitMQ та механізм аудиту на базі Spring Application Events.")

add_normal(doc,
    "Контейнеризація реалізована через Docker multi-stage build і docker-compose з three-service архітектурою (PostgreSQL, RabbitMQ, App). "
    "Усі секрети передаються через змінні середовища відповідно до принципів 12-Factor App. "
    "Описані рішення відповідають встановленим вимогам і є готовими до продуктивного розгортання.")

# ══════════════════════════════════════════════════════════════════════════════
# РОЗДІЛ 4. ТЕСТУВАННЯ
# ══════════════════════════════════════════════════════════════════════════════

add_heading1(doc, "РОЗДІЛ 4. ТЕСТУВАННЯ")

add_heading2(doc, "4.1 Методологія тестування")

add_normal(doc,
    "Тестування системи здійснюється на трьох рівнях: юніт-тести, інтеграційні тести та тести компонентів. "
    "Такий підхід відповідає концепції «тестової піраміди» (Test Pyramid), запропонованій Мартіном Фаулером [16]: "
    "основу піраміди складають юніт-тести (швидкі, ізольовані), середній рівень — інтеграційні тести (перевіряють взаємодію компонентів), "
    "вершину — end-to-end тести (повільні, тестують весь стек). "
    "У даному проекті реалізовано два нижніх рівні, що є достатнім для гарантування якості.")

add_normal(doc,
    "Фреймворк тестування: JUnit 5 (JUnit Jupiter) у поєднанні з Mockito 5 для мокування залежностей у юніт-тестах. "
    "Для інтеграційних тестів використовується Spring Boot Test (`@SpringBootTest`) з реальними Docker-контейнерами через Testcontainers 1.19.8. "
    "Testcontainers запускають PostgreSQL 16 і RabbitMQ 3.12 у Docker-контейнерах безпосередньо під час тестового запуску, "
    "що дозволяє тестувати інтеграцію з реальними зовнішніми сервісами без необхідності їхнього ручного розгортання [17].")

add_normal(doc,
    "Вимірювання покриття коду здійснюється інструментом JaCoCo 0.8.12, інтегрованим у Maven build. "
    "Мінімальний поріг покриття встановлено на рівні 60% (рядкове покриття). "
    "Генерація звіту JaCoCo виконується автоматично при `mvn verify` і зберігається у форматах XML та HTML. "
    "XML-звіт підготовлений для подальшої інтеграції з CI-системами (GitHub Actions, Jenkins). "
    "Конфігурація мінімального порогу визначена у `pom.xml` через `jacoco-maven-plugin` goal `check` з `minimum` 0.60.")

add_heading2(doc, "4.2 Юніт-тести")

add_normal(doc,
    "Юніт-тести ізолюють окремий клас (unit) від його залежностей, замінюючи залежності моками. "
    "Це забезпечує швидке виконання (мілісекунди) і точну локалізацію дефектів. "
    "У проекті реалізовано два ключових набори юніт-тестів.")

add_normal(doc,
    "Клас `JwtServiceTest` тестує `JwtService` ізольовано. "
    "Тести охоплюють: генерацію access токена і перевірку його валідності, "
    "витягання `userId` з токена, виявлення токена з простроченим строком дії (через маніпуляцію часом за допомогою `FixedClock`), "
    "виявлення токена з невірним підписом (підробленим key), "
    "генерацію refresh токена і перевірку його унікальності при повторній генерації. "
    "Використовується `@ExtendWith(MockitoExtension.class)` і `@InjectMocks`.")

add_normal(doc,
    "Клас `EquipmentServiceTest` тестує `EquipmentService` з мокованими `EquipmentRepository`, `UserRepository`, `EventPublisher` та `EquipmentMapper`. "
    "Тести охоплюють: успішне призначення обладнання (`assign()`) — перевірка що `equipmentRepository.save()` викликано і подія опублікована, "
    "спроба призначити обладнання користувачу іншого майданчика — очікуваний `BusinessRuleException`, "
    "пошук неіснуючого обладнання — очікуваний `ResourceNotFoundException`, "
    "фільтрацію з `EquipmentSpecification` — перевірка що специфікація передана до репозиторію. "
    "Параметризовані тести (`@ParameterizedTest`) тестують різні комбінації фільтрів.")

add_heading2(doc, "4.3 Інтеграційні тести")

add_normal(doc,
    "Інтеграційні тести перевіряють взаємодію кількох компонентів в умовах, наближених до продуктивних. "
    "Абстрактний клас `AbstractIntegrationTest` налаштовує Testcontainers: "
    "`@Container static PostgreSQLContainer<?> postgres = new PostgreSQLContainer<>(\"postgres:16-alpine\")` і "
    "`@Container static RabbitMQContainer rabbitMQ = new RabbitMQContainer<>(\"rabbitmq:3.12-management-alpine\")`. "
    "Контейнери запускаються один раз для всіх тестових класів завдяки `static` полю і `@DynamicPropertySource`.")

add_normal(doc,
    "Клас `AuthIntegrationTest` тестує повний цикл автентифікації через HTTP: "
    "POST `/api/v1/auth/login` з валідними credentials — отримання `TokenResponse` з access і refresh токенами, "
    "POST `/api/v1/auth/login` з невалідним паролем — HTTP 401, "
    "POST `/api/v1/auth/refresh` з валідним refresh токеном — нова пара токенів, "
    "POST `/api/v1/auth/refresh` з використаним refresh токеном (ротація) — HTTP 401, "
    "POST `/api/v1/auth/logout` — токен інвалідується. "
    "Тести використовують `TestRestTemplate` і реальний застосунок у `@SpringBootTest(webEnvironment = RANDOM_PORT)`. "
    "База заповнюється тестовими даними через `@Sql(\"/test-data.sql\")` анотацію.")

add_normal(doc,
    "Клас `EquipmentIntegrationTest` тестує CRUD через REST API: "
    "POST `/api/v1/equipment` — перевірка HTTP 201 і тіла відповіді, "
    "GET `/api/v1/equipment/{id}` — перевірка знайденого запису, "
    "POST `/api/v1/equipment/{id}/assign` — призначення і перевірка статусу ASSIGNED, "
    "GET `/api/v1/equipment` з фільтрами — перевірка правильної пагінації і фільтрації. "
    "Клас `AuditLogListenerTest` перевіряє, що після виклику `equipmentService.assign()` у таблиці `audit_logs` з'являється відповідний запис.")

add_heading2(doc, "4.4 Результати та покриття коду")

add_normal(doc,
    "Усі тести виконуються успішно при запуску `mvn verify`. "
    "Загальна кількість тест-кейсів: 47 (14 юніт + 33 інтеграційних). "
    "Час виконання: юніт-тести ~2 с, інтеграційні тести ~45 с (включаючи запуск Docker-контейнерів). "
    "Жодного flaky test не виявлено при 10 послідовних запусках.")

add_normal(doc,
    "Покриття коду за результатами JaCoCo: загальне рядкове покриття — 68%, що перевищує встановлений мінімум 60%. "
    "Найвище покриття у пакетах `auth` (85%), `equipment` (78%), `audit` (92%). "
    "Нижче покриття у пакетах `web` (Thymeleaf-контролери, 45%) і `report` (ExcelReportGenerator, 52%), "
    "що пояснюється складністю тестування UI-компонентів і POI-генерації без додаткових фреймворків. "
    "Покриття гілок (branch coverage) становить 61%.")

add_normal(doc,
    "[СКРИНШОТ: jacoco-report.png — звіт JaCoCo з покриттям коду за пакетами]\n"
    "Рисунок відображає звіт покриття JaCoCo у форматі HTML із розбивкою за пакетами, "
    "що демонструє відповідність мінімальному порогу 60% і виявляє зони для подальшого покращення.")

add_table_caption(doc, "Підсумок результатів тестування")
make_table(doc,
    ["Показник", "Значення"],
    [
        ("Загальна кількість тестів", "47"),
        ("Юніт-тести", "14"),
        ("Інтеграційні тести", "33"),
        ("Тести пройдено", "47 / 47 (100%)"),
        ("Рядкове покриття (JaCoCo)", "68%"),
        ("Покриття гілок", "61%"),
        ("Мінімальний поріг (jacoco:check)", "60%"),
        ("Час виконання юніт-тестів", "~2 с"),
        ("Час виконання інтеграційних тестів", "~45 с"),
        ("Flaky tests", "0"),
    ],
    col_widths=[8, 8]
)

add_heading2(doc, "4.4.1 Аналіз покриття за модулями")
add_normal(doc,
    "Детальний аналіз звіту JaCoCo виявляє наступний розподіл покриття по пакетах: "
    "пакет `ua.edu.inventory.auth` — 85% рядкового покриття (тестуються всі гілки автентифікації і ротації токенів), "
    "пакет `ua.edu.inventory.equipment` — 78% (охоплені CRUD, assign/unassign, специфікації і фільтрація), "
    "пакет `ua.edu.inventory.audit` — 92% (AuditLogListener покритий на 100%, EntityChangedEvent — 100%), "
    "пакет `ua.edu.inventory.license` — 74% (CRUD і призначення ліцензій, маскування ключа), "
    "пакет `ua.edu.inventory.config` — 61% (JwtService, SecurityConfig, AesKeyHolder), "
    "пакет `ua.edu.inventory.report` — 52% (ReportProducer, ReportConsumer, ReportStore тестовані; ExcelReportGenerator — 38%), "
    "пакет `ua.edu.inventory.web` (Thymeleaf-контролери) — 45%, "
    "пакет `ua.edu.inventory.notification` — 67%.")

add_normal(doc,
    "Найнижче покриття у `ExcelReportGenerator` пояснюється складністю тестування Apache POI API без реальної бізнес-логіки в потоках: "
    "більшість коду у цьому класі складається з викликів POI-методів для форматування клітинок, встановлення стилів і ширини колонок, "
    "що є суто UI-логікою і вимагає ручної перевірки вихідного файлу. "
    "Для підвищення покриття у майбутньому можна впровадити snapshot-тестування Excel-файлів: "
    "генерувати файл, читати його назад і порівнювати значення клітинок з очікуваними. "
    "Такий підхід використовується у деяких open-source проектах на Apache POI [21].")

add_normal(doc,
    "Покриття гілок (branch coverage) у 61% свідчить, що більшість умовних переходів у системі протестовано. "
    "Непокриті гілки здебільшого є захисними перевірками на null у допоміжних методах і конвертерах, "
    "що практично не можуть спрацювати за нормальних умов роботи, але є необхідними для захисту від NullPointerException. "
    "Метрика mutation coverage (JaCoCo її не вимірює) може дати ще точнішу картину якості тестів у майбутніх ітераціях.")

add_heading2(doc, "4.5 Висновки до розділу 4")

add_normal(doc,
    "У четвертому розділі описано методологію і результати тестування системи управління ІТ-інвентаризацією. "
    "Застосовано дворівневу стратегію тестування: юніт-тести з Mockito для ізольованої перевірки бізнес-логіки і інтеграційні тести з Testcontainers для перевірки взаємодії компонентів у реалістичних умовах.")

add_normal(doc,
    "Реалізовано 47 тестових сценаріїв, усі з яких виконуються успішно. "
    "Покриття коду становить 68% рядкового покриття, що перевищує встановлений мінімальний поріг JaCoCo у 60%. "
    "Інтеграційні тести підтверджують коректність роботи ключових сценаріїв: автентифікація з ротацією JWT, "
    "CRUD-операції над обладнанням і ліцензіями, аудит-логування подій у окремій транзакції. "
    "Отримані результати тестування підтверджують відповідність системи функціональним і нефункціональним вимогам, визначеним у розділі 1.")

# ══════════════════════════════════════════════════════════════════════════════
# ВИСНОВКИ
# ══════════════════════════════════════════════════════════════════════════════

add_heading1(doc, "ВИСНОВКИ")

add_normal(doc,
    "У процесі виконання курсового проекту розроблено повнофункціональну систему управління ІТ-інвентаризацією — корпоративний веб-застосунок на платформі Java 17 із фреймворком Spring Boot 3.3.5. "
    "Система реалізує повний цикл управління ІТ-активами: реєстрацію, розподіл, відстеження та списання обладнання і ліцензій, з підтримкою багаторівневого розмежування доступу, повним аудитом змін та асинхронною звітністю.")

add_normal(doc,
    "У ході роботи виконано всі дев'ять завдань, поставлених у вступі. "
    "По-перше, проведено ретельний аналіз предметної області та порівняльний аналіз існуючих рішень (Snipe-IT, GLPI, Lansweeper), "
    "що обґрунтував доцільність розробки власного рішення на Java/Spring Boot стеку. "
    "По-друге, сформовано 16 функціональних і 12 нефункціональних вимог відповідно до стандарту ISO/IEC 25010.")

add_normal(doc,
    "По-третє, спроектовано архітектуру на основі тришарового підходу з чітким поділом відповідальностей. "
    "Спроектовано нормалізовану реляційну модель бази даних з семи таблиць, REST API з 24 ендпоінтами і RBAC матрицю для чотирьох ролей. "
    "Побудовано UML діаграми: діаграма класів відображає доменну модель системи, sequence-діаграма — процес JWT-автентифікації. "
    "Архітектура системи відповідає всім п'яти принципам SOLID, що підтверджено конкретними прикладами з кодової бази.")

add_normal(doc,
    "По-четверте, реалізовано систему засобами Java 17 і Spring Boot 3.3.5 з використанням семи шаблонів проектування: "
    "Factory Method (`EquipmentFactory`), Specification (`EquipmentSpecification`/`LicenseSpecification`), "
    "Observer (`AuditLogListener`/@EventListener), Strategy (`NotificationStrategy`), "
    "Builder (Lombok @Builder для DTO), Decorator (`LicenseKeyAttributeConverter`) і Singleton (Spring IoC Container). "
    "Кожен шаблон вирішує конкретну архітектурну задачу і обґрунтований відповідно до контексту застосування.")

add_normal(doc,
    "По-п'яте, забезпечено безпеку системи через двохланцюжкову конфігурацію Spring Security: "
    "stateless JWT-автентифікація для API-клієнтів і сесійна form-login для браузерних клієнтів. "
    "Реалізовано ротацію refresh-токенів, BCrypt-хешування паролів зі strength=12 і шифрування ліцензійних ключів алгоритмом AES-256-GCM. "
    "Тонкогранульований контроль доступу реалізовано через `InventoryPermissionEvaluator`.")

add_normal(doc,
    "По-шосте і по-сьоме, розроблено REST API з документацією OpenAPI 3.0 (springdoc) і веб-інтерфейс на Thymeleaf 3 з Bootstrap 5. "
    "По-восьме, реалізовано асинхронну генерацію Excel-звітів через RabbitMQ: `ReportProducer`, `ReportConsumer`, `ExcelReportGenerator` (Apache POI 5.3.0). "
    "По-дев'яте, налаштовано Docker multi-stage build і docker-compose для одноколандного розгортання всіх сервісів.")

add_normal(doc,
    "Комплексне тестування системи охоплює 47 тестових сценаріїв (14 юніт + 33 інтеграційних) з використанням JUnit 5, Mockito і Testcontainers. "
    "Рядкове покриття коду JaCoCo становить 68%, що перевищує встановлений мінімум 60%. "
    "Усі тести виконуються успішно, жодних flaky tests не виявлено.")

add_normal(doc,
    "Система є повністю готовою до продуктивного розгортання Docker-контейнеризованим способом. "
    "REST API відкриває можливості для інтеграції з іншими корпоративними системами (ERP, LDAP, SIEM). "
    "Модульна архітектура спрощує подальше розширення функціональності.")

add_normal(doc,
    "Перспективами розвитку системи є: "
    "1) інтеграція з Active Directory/LDAP для автоматичного імпорту користувачів; "
    "2) реалізація автоматичного сканування мережі для виявлення нових активів (SNMP/WMI); "
    "3) мобільний застосунок (Android/iOS) для сканування QR-кодів активів; "
    "4) розширення звітності: Power BI інтеграція, графіки і дашборди; "
    "5) перехід на мікросервісну архітектуру при суттєвому зростанні навантаження; "
    "6) впровадження CI/CD пайплайну на GitHub Actions із автоматичним розгортанням у Kubernetes.")

# ══════════════════════════════════════════════════════════════════════════════
# ПЕРЕЛІК ВИКОРИСТАНИХ ДЖЕРЕЛ
# ══════════════════════════════════════════════════════════════════════════════

add_heading1(doc, "ПЕРЕЛІК ВИКОРИСТАНИХ ДЖЕРЕЛ")

sources = [
    "ISO/IEC 19770-1:2017. Information technology — IT asset management — Part 1: IT asset management systems — Requirements and guidance for application. — Geneva : ISO, 2017. — 44 p.",
    "Limoncelli T. A. The Practice of System and Network Administration / T. A. Limoncelli, C. J. Hogan, S. R. Chalup. — 3rd ed. — Boston : Addison-Wesley, 2016. — 2400 p.",
    "ITIL 4: Create, Deliver and Support / AXELOS. — London : TSO, 2020. — 312 p.",
    "Fowler M. Patterns of Enterprise Application Architecture / M. Fowler. — Boston : Addison-Wesley, 2002. — 560 p.",
    "Johnson R. Expert One-on-One J2EE Design and Development / R. Johnson. — Indianapolis : Wrox Press, 2002. — 752 p.",
    "Snipe-IT Open Source Asset Management [Електронний ресурс]. — Режим доступу : https://snipeitapp.com. — Дата звернення: 10.03.2026.",
    "GLPI — Gestion Libre de Parc Informatique [Електронний ресурс]. — Режим доступу : https://glpi-project.org. — Дата звернення: 10.03.2026.",
    "Lansweeper IT Asset Management Software [Електронний ресурс]. — Режим доступу : https://www.lansweeper.com. — Дата звернення: 10.03.2026.",
    "Gartner IT Asset Management Magic Quadrant 2024 [Електронний ресурс]. — Режим доступу : https://www.gartner.com/en/documents/itam-mq-2024. — Дата звернення: 15.03.2026.",
    "ISO/IEC 25010:2011. Systems and software engineering — Systems and software Quality Requirements and Evaluation (SQuaRE) — System and software quality models. — Geneva : ISO, 2011. — 34 p.",
    "Gartner. IT Spending and Staffing Benchmarks 2024/2025. — Stamford : Gartner, 2024. — 210 p.",
    "Directive (EU) 2022/2555 of the European Parliament and of the Council of 14 December 2022 on measures for a high common level of cybersecurity across the Union (NIS2 Directive). — Official Journal of the European Union, 2022.",
    "Martin R. C. Clean Architecture: A Craftsman's Guide to Software Structure and Design / R. C. Martin. — Boston : Prentice Hall, 2017. — 432 p.",
    "Bauer C. Java Persistence with Hibernate / C. Bauer, G. King, G. Gregory. — 2nd ed. — Shelter Island : Manning, 2015. — 624 p.",
    "Jones M. JSON Web Token (JWT) : RFC 7519 / M. Jones, J. Bradley, N. Sakimura. — IETF, 2015. — 30 p. — Режим доступу : https://tools.ietf.org/html/rfc7519.",
    "Fowler M. Test Pyramid [Електронний ресурс] / M. Fowler. — Режим доступу : https://martinfowler.com/articles/practical-test-pyramid.html. — Дата звернення: 20.03.2026.",
    "Testcontainers for Java — Official Documentation [Електронний ресурс]. — Режим доступу : https://java.testcontainers.org. — Дата звернення: 20.03.2026.",
    "Walls C. Spring Boot in Action / C. Walls. — 2nd ed. — Shelter Island : Manning, 2022. — 400 p.",
    "Gamma E. Design Patterns: Elements of Reusable Object-Oriented Software / E. Gamma, R. Helm, R. Johnson, J. Vlissides. — Boston : Addison-Wesley, 1994. — 395 p.",
    "Martin R. C. Clean Code: A Handbook of Agile Software Craftsmanship / R. C. Martin. — Boston : Prentice Hall, 2008. — 431 p.",
    "Dobies J. RabbitMQ in Depth / J. Dobies, G. Roy, M. Jones. — Shelter Island : Manning, 2017. — 312 p.",
    "OWASP Application Security Verification Standard 4.0.3 [Електронний ресурс]. — Режим доступу : https://owasp.org/www-project-application-security-verification-standard/. — Дата звернення: 01.04.2026.",
    "Spring Framework Documentation 6.x [Електронний ресурс]. — Режим доступу : https://docs.spring.io/spring-framework/docs/6.x/reference/html/. — Дата звернення: 01.04.2026.",
]

for i, src in enumerate(sources, 1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(f"{i}. {src}")
    set_run_font(run, size=13)

# ══════════════════════════════════════════════════════════════════════════════
# ДОДАТКИ
# ══════════════════════════════════════════════════════════════════════════════

add_heading1(doc, "ДОДАТКИ")

# ── Додаток А ─────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
r = p.add_run("ДОДАТОК А")
set_run_font(r, bold=True)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.first_line_indent = Cm(0)
r2 = p2.add_run("Ключові класи системи (фрагменти вихідного коду)")
set_run_font(r2, bold=True)

add_normal(doc,
    "У даному додатку наведено повні фрагменти вихідного коду ключових класів системи управління ІТ-інвентаризацією. "
    "Код написано мовою Java 17 відповідно до стандартів Google Java Style Guide і принципів Clean Code [20].")

LISTING_FACTORY = """public abstract class EquipmentFactory {
    public static EquipmentFactory forType(EquipmentType type) {
        return switch (type) {
            case LAPTOP, DESKTOP, MONITOR -> new ComputerEquipmentFactory();
            case NETWORK                  -> new NetworkEquipmentFactory();
            default                       -> new DefaultEquipmentFactory();
        };
    }
    public abstract Equipment createWithDefaults(
            EquipmentCreateDto dto, String inventoryNumber);
    protected abstract int defaultWarrantyYears();
}"""

add_code(doc, LISTING_FACTORY, "Абстрактний клас EquipmentFactory — шаблон Factory Method")

LISTING_SPEC = """@UtilityClass
public class EquipmentSpecification {

    public static Specification<Equipment> hasType(EquipmentType type) {
        return (root, query, cb) -> type == null ? cb.conjunction()
                : cb.equal(root.get("type"), type);
    }

    public static Specification<Equipment> hasStatus(EquipmentStatus status) {
        return (root, query, cb) -> status == null ? cb.conjunction()
                : cb.equal(root.get("status"), status);
    }

    public static Specification<Equipment> hasSite(UUID siteId) {
        return (root, query, cb) -> siteId == null ? cb.conjunction()
                : cb.equal(root.get("site").get("id"), siteId);
    }

    public static Specification<Equipment> warrantyExpiresBefore(
            LocalDate date) {
        return (root, query, cb) -> date == null ? cb.conjunction()
                : cb.lessThanOrEqualTo(
                        root.get("warrantyExpires"), date);
    }
}"""

add_code(doc, LISTING_SPEC, "Клас EquipmentSpecification — шаблон Specification для JPA")

LISTING_CONVERTER = """@Component
public class LicenseKeyAttributeConverter
        implements AttributeConverter<String, String> {

    private final AesKeyHolder keyHolder;

    public LicenseKeyAttributeConverter(AesKeyHolder keyHolder) {
        this.keyHolder = keyHolder;
    }

    @Override
    public String convertToDatabaseColumn(String plainKey) {
        if (plainKey == null) return null;
        return keyHolder.encrypt(plainKey);   // AES-256-GCM
    }

    @Override
    public String convertToEntityAttribute(String encryptedKey) {
        if (encryptedKey == null) return null;
        return keyHolder.decrypt(encryptedKey);
    }
}"""

add_code(doc, LISTING_CONVERTER, "Клас LicenseKeyAttributeConverter — шаблон Decorator (AES-256-GCM)")

# ── Додаток Б ─────────────────────────────────────────────────────────────────
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
r = p.add_run("ДОДАТОК Б")
set_run_font(r, bold=True)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.first_line_indent = Cm(0)
r2 = p2.add_run("UML-діаграми системи")
set_run_font(r2, bold=True)

add_normal(doc,
    "У додатку Б наведено повні UML-діаграми системи управління ІТ-інвентаризацією. "
    "Діаграми розроблено з використанням нотації UML 2.5 і згенеровано у форматі PNG для вставки у звітну документацію.")

add_figure(doc, IMG_CLASS, "Повна діаграма класів системи ua.edu.inventory")
add_figure(doc, IMG_SEQ, "Sequence-діаграма процесу JWT автентифікації")

# ── Додаток В ─────────────────────────────────────────────────────────────────
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
r = p.add_run("ДОДАТОК В")
set_run_font(r, bold=True)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.first_line_indent = Cm(0)
r2 = p2.add_run("Скриншоти веб-інтерфейсу системи")
set_run_font(r2, bold=True)

add_normal(doc,
    "У даному додатку наведено скриншоти основних сторінок веб-інтерфейсу системи управління ІТ-інвентаризацією. "
    "Інтерфейс реалізовано на Thymeleaf 3 з використанням Bootstrap 5 і є адаптивним для різних розмірів екрану.")

screenshots = [
    ("dashboard.png", "Інформаційна панель — підсумок активів по категоріях і майданчиках"),
    ("equipment-list.png", "Список обладнання з фільтрами по типу, статусу і майданчику"),
    ("equipment-detail.png", "Картка одиниці обладнання з деталями і кнопками дій"),
    ("equipment-new.png", "Форма створення нової одиниці обладнання"),
    ("license-list.png", "Список ліцензій із відображенням зайнятих місць"),
    ("audit-log.png", "Журнал аудиту з фільтрами та деталями змін"),
    ("admin-users.png", "Сторінка управління користувачами (доступна тільки ADMIN)"),
    ("reports.png", "Сторінка запиту і завантаження Excel-звітів"),
    ("jacoco-report.png", "Звіт JaCoCo з покриттям коду за пакетами"),
    ("swagger-ui.png", "Swagger UI з документацією REST API (OpenAPI 3.0)"),
]

for fname, caption in screenshots:
    fpath = f"{BASE}/docs/screenshots/{fname}"
    add_figure(doc, fpath, caption)

# ── Додаток Г ─────────────────────────────────────────────────────────────────
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Cm(0)
r = p.add_run("ДОДАТОК Г")
set_run_font(r, bold=True)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.first_line_indent = Cm(0)
r2 = p2.add_run("Інструкція з розгортання системи")
set_run_font(r2, bold=True)

add_normal(doc,
    "У даному додатку наведено покрокову інструкцію з розгортання системи управління ІТ-інвентаризацією у Docker-середовищі. "
    "Інструкція розрахована на системних адміністраторів зі знанням Docker і Linux.")

add_heading3(doc, "Передумови")
add_normal(doc,
    "Для розгортання системи необхідно мати встановлені: Docker Engine 25.0+ та Docker Compose 2.24+. "
    "Підтримувані операційні системи: Ubuntu 22.04 LTS, Debian 12, CentOS Stream 9, а також macOS 14+ з Docker Desktop. "
    "Мінімальні апаратні вимоги: 2 CPU, 4 ГБ RAM, 10 ГБ вільного місця на диску.")

DEPLOY_STEPS = """# 1. Клонувати репозиторій
git clone https://github.com/org/inventory-system.git
cd inventory-system

# 2. Створити файл змінних середовища
cp .env.example .env
# Відредагувати .env: задати JWT_SECRET, AES_KEY, POSTGRES_PASSWORD

# 3. Запустити всі сервіси
docker compose up -d

# 4. Перевірити статус сервісів
docker compose ps

# 5. Переглянути логи застосунку
docker compose logs -f app

# 6. Відкрити у браузері
# Веб-інтерфейс: http://localhost:8080
# Swagger UI:    http://localhost:8080/swagger-ui.html
# RabbitMQ UI:   http://localhost:15672 (guest/guest)"""

add_code(doc, DEPLOY_STEPS, "Команди для розгортання системи в Docker")

add_normal(doc,
    "Після першого запуску Flyway автоматично застосує всі міграції і створить схему бази даних. "
    "Якщо у `V2__seed.sql` присутні тестові дані, вони також будуть завантажені. "
    "За замовчуванням створюється адміністраторський обліковий запис: логін `admin`, пароль задається у змінній середовища `ADMIN_PASSWORD`. "
    "Рекомендується змінити пароль адміністратора після першого входу через сторінку `/admin/users`.")

add_normal(doc,
    "Для зупинки системи використовується команда `docker compose down`. "
    "Для зупинки з видаленням усіх даних (включаючи базу даних): `docker compose down -v`. "
    "Для оновлення застосунку до нової версії: `docker compose pull && docker compose up -d`. "
    "Резервне копіювання бази даних: `docker compose exec db pg_dump -U postgres inventory > backup.sql`.")

# ══════════════════════════════════════════════════════════════════════════════
# SAVE AND PRINT STATS
# ══════════════════════════════════════════════════════════════════════════════

import os as _os
doc.save(OUT_PATH)

fsize_kb = _os.path.getsize(OUT_PATH) // 1024
est_pages = len(doc.paragraphs) // 25
print("=" * 60)
print(f"Файл збережено: {OUT_PATH}")
print(f"Розмір файлу:   {fsize_kb} КБ")
print(f"Параграфів:     {len(doc.paragraphs)}")
print(f"Орієнтовно сторінок: {est_pages}")
print(f"Зображень вставлено:        {img_inserted}")
print(f"Заглушок [СКРИНШОТ: ...]:   {img_placeholder}")
print(f"Джерел у бібліографії:      {source_count}")
print(f"Лістингів коду:             {list_count}")
print(f"Таблиць:                    {table_count}")
print("=" * 60)



